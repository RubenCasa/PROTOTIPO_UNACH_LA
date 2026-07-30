import docx
from docx.shared import Pt, Inches
import sys

try:
    doc = docx.Document('INFORME_FINAL_COMPLETO_UNACH_LA_Actualizado.docx')
except Exception as e:
    print(f"Error abriendo documento: {e}")
    sys.exit(1)

# Buscar "ÍNDICE GENERAL" (puede tener problemas de codificación, busquemos "NDICE GENERAL")
insert_idx = -1
for i, p in enumerate(doc.paragraphs):
    if "NDICE GENERAL" in p.text.upper():
        insert_idx = i
        break

if insert_idx == -1:
    print("No se encontró el Índice General, insertando al inicio después de los títulos.")
    insert_idx = 19 # Un valor seguro asumiendo la portada

target_paragraph = doc.paragraphs[insert_idx]

# Insertar párrafos ANTES del Índice
p1 = target_paragraph.insert_paragraph_before('Resumen de Fases del Cronograma (Overview)')
p1.style = 'Heading 1'

target_paragraph.insert_paragraph_before('A continuación, se presenta un resumen de las tres fases principales que estructuraron este proyecto de ayudantía de investigación:')

target_paragraph.insert_paragraph_before('Fase 1: Alineación Normativa y Fundamentación (Semanas 1-5): Enfocada en el diagnóstico tecnológico, análisis legal de protección de datos y lineamientos éticos.', style='List Bullet')
target_paragraph.insert_paragraph_before('Fase 2: Diseño e Implementación ML (Semanas 7-11): Enfocada en la preparación de datos académicos, desarrollo y evaluación de los modelos predictivos, y diseño de KPIs.', style='List Bullet')
target_paragraph.insert_paragraph_before('Fase 3: Documentación y Cierre (Semanas 12-16): Enfocada en la consolidación del prototipo funcional, elaboración del informe final y entrega de resultados.', style='List Bullet')

target_paragraph.insert_paragraph_before('') # Espacio

doc.save('INFORME_FINAL_COMPLETO_UNACH_LA_Actualizado.docx')
print("Resumen de fases insertado correctamente al inicio del documento.")
