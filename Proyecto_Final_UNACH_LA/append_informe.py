import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import sys

try:
    doc = docx.Document('INFORME_FINAL_COMPLETO_UNACH_LA.docx')
except Exception as e:
    print(f"Error abriendo documento: {e}")
    sys.exit(1)

# Función auxiliar para agregar título
def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return h

doc.add_page_break()

add_heading(doc, 'Cronograma de Actividades y Resultados Alcanzados', 1)
doc.add_paragraph('A continuación se detalla el cronograma de actividades ejecutado durante el proyecto, junto con la descripción de los resultados obtenidos en cada una de las fases y entregables planificados.')

add_heading(doc, '1. Cronograma de Actividades', 2)

table_data = [
    ["Nº", "Fase / Objetivo", "Actividades principales", "Entregables / Productos", "Fecha de Ejecución", "Horas"],
    ["1", "Fase 1: Alineación Normativa y Fundamentación", "Revisión de normativa institucional UNACH y Ley Orgánica de Protección de Datos Personales", "Matriz normativa y análisis legal", "27 Abr - 3 May", "8h"],
    ["2", "Resultado de Aprendizaje 1", "Investigación de casos de éxito de ML y revisión de literatura de Learning Analytics", "Estado del arte y bibliografía sistematizada", "4 May - 10 May", "8h"],
    ["3", "Diagnóstico Tecnológico", "Caracterización del ecosistema tecnológico UNACH y análisis de plataformas SICO/Moodle", "Informe de ecosistema y flujos de datos", "11 May - 17 May", "8h"],
    ["4", "Gobierno y Ética de Datos", "Definición de criterios éticos, privacidad y gobernanza institucional", "Documento de lineamientos éticos y de gobernanza", "18 May - 24 May", "6h"],
    ["5", "Consolidación de Fase 1", "Llenado de registros de ayudantía y elaboración de informe de diagnóstico", "Informe parcial académico", "25 May - 31 May", "6h"],
    ["6", "SEMANA DE EXÁMENES", "No se programan actividades", "—", "1 Jun - 7 Jun", "0h"],
    ["7", "Fase 2: Diseño e Implementación ML", "Preparación y limpieza de datos académicos", "Dataset procesado y documentado", "8 Jun - 14 Jun", "10h"],
    ["8", "Resultado de Aprendizaje 2", "Diseño de modelos ML para diagnóstico académico", "Arquitectura y diseño metodológico de modelos", "15 Jun - 21 Jun", "10h"],
    ["9", "Desarrollo Técnico", "Implementación de modelos ML en Python y scikit-learn", "Código funcional y scripts de análisis", "22 Jun - 28 Jun", "10h"],
    ["10", "Evaluación de Modelos", "Validación y pruebas de rendimiento (AUC, precisión, recall)", "Resultados de métricas y validación", "29 Jun - 5 Jul", "8h"],
    ["11", "Visualización y KPIs", "Diseño de dashboards e indicadores clave de desempeño", "Propuesta de KPIs y visualizaciones", "6 Jul - 12 Jul", "6h"],
    ["12", "Fase 3: Documentación y Cierre", "Consolidación del prototipo funcional y modelo UNACH-LA", "Prototipo funcional completo", "13 Jul - 19 Jul", "8h"],
    ["13", "Resultado de Aprendizaje 3", "Elaboración del informe académico final y validación de resultados", "Informe final académico", "20 Jul - 26 Jul", "6h"],
    ["14", "Cierre Administrativo", "Llenado final de registros y formatos de ayudantía", "Registros completos y firmados", "20 Jul - 26 Jul", "1h"],
    ["15", "SEMANA DE EXÁMENES", "No se programan actividades", "—", "27 Jul - 2 Ago", "0h"],
    ["16", "Sustentación y Entrega Final", "Presentación de informe y entrega de productos técnicos", "Sustentación, código y manuales entregados", "3 Ago - 7 Ago", "1h"],
    ["", "Total", "", "", "", "96h"]
]

table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
for i, heading in enumerate(table_data[0]):
    hdr_cells[i].text = heading

for row in table_data[1:]:
    row_cells = table.add_row().cells
    for i, text in enumerate(row):
        row_cells[i].text = text

add_heading(doc, '2. Explicación de Resultados por Fase y Actividad', 2)

resultados = [
    ("Fase 1: Alineación Normativa y Fundamentación", "Resultados (Matriz normativa y análisis legal): Se logró establecer un marco jurídico sólido para el manejo de datos estudiantiles. El análisis garantizó que el procesamiento de la información académica (notas, asistencias) para los modelos predictivos cumple estrictamente con la normativa de la UNACH y la Ley Orgánica de Protección de Datos Personales, mitigando riesgos legales y asegurando la confidencialidad de la información de los estudiantes."),
    ("Resultado de Aprendizaje 1", "Resultados (Estado del arte y bibliografía): Se construyó una base teórica y práctica sólida mediante la revisión exhaustiva de casos de éxito en Learning Analytics. Esto permitió identificar los algoritmos de Machine Learning más efectivos (como Random Forest, XGBoost y Regresión Logística) utilizados en otras instituciones de educación superior para la predicción de rendimiento y deserción, sirviendo como fundamento para las decisiones arquitectónicas del proyecto."),
    ("Diagnóstico Tecnológico", "Resultados (Informe de ecosistema y flujos de datos): Se mapeó exitosamente la infraestructura tecnológica actual (plataformas SICO y Moodle). Se identificaron las fuentes de datos clave, los formatos de exportación y las limitaciones del sistema actual. Este diagnóstico permitió diseñar un flujo de datos (pipeline) viable para la extracción, transformación y carga (ETL) hacia el nuevo sistema predictivo."),
    ("Gobierno y Ética de Datos", "Resultados (Documento de lineamientos éticos): Se definieron reglas claras sobre quién tiene acceso a las predicciones y cómo se deben comunicar estos resultados a los estudiantes (evitando sesgos y estigmatización). Se establecieron protocolos para anonimizar los datos durante la fase de entrenamiento de los modelos, garantizando un uso ético y responsable de la inteligencia artificial en el contexto universitario."),
    ("Fase 2: Diseño e Implementación ML", "Resultados (Dataset procesado y documentado): Se generó un dataset estructurado, limpio y listo para el entrenamiento de modelos. Se resolvieron problemas de datos faltantes, se normalizaron variables (como calificaciones y porcentajes de asistencia) y se documentó el diccionario de datos, lo cual es fundamental para asegurar que los modelos aprendan patrones reales y no ruido."),
    ("Resultado de Aprendizaje 2", "Resultados (Arquitectura y diseño metodológico): Se definió la arquitectura del sistema, separando claramente el backend (API en Python/FastAPI) y el frontend (Dashboard en React). Metodológicamente, se determinaron las variables predictoras clave y la variable objetivo (riesgo académico), estableciendo cómo los modelos de ML integrarían la información histórica para generar alertas tempranas."),
    ("Desarrollo Técnico", "Resultados (Código funcional y scripts de análisis): Se programaron y entrenaron los modelos de predicción utilizando Python y Scikit-Learn. El resultado fue la creación de un motor analítico funcional capaz de recibir datos de un estudiante e inferir su probabilidad de éxito o riesgo académico, empaquetado en una API robusta y lista para ser consumida por la interfaz de usuario."),
    ("Evaluación de Modelos", "Resultados (Métricas y validación): Los modelos fueron evaluados con rigor científico utilizando métricas como el Área Bajo la Curva (AUC), Precisión (Precision) y Exhaustividad (Recall). Los resultados demostraron un nivel de precisión satisfactorio, asegurando que el modelo es confiable identificando a estudiantes en riesgo (reduciendo falsos positivos y falsos negativos) antes de integrarlo al prototipo final."),
    ("Visualización y KPIs", "Resultados (Propuesta de KPIs y visualizaciones): Se diseñó una interfaz interactiva y amigable orientada a los docentes y autoridades. Los resultados se materializaron en paneles de control (dashboards) que traducen las probabilidades del modelo de ML en indicadores clave de desempeño (KPIs) visuales, alertas gráficas y recomendaciones automatizadas, facilitando la toma de decisiones pedagógicas."),
    ("Fase 3: Documentación y Cierre (y Resultado de Aprendizaje 3)", "Resultados (Prototipo y validación de resultados): Se logró consolidar todas las piezas (modelos, API, dashboard) en un Prototipo Funcional Completo denominado \"Modelo UNACH-LA\". Se validó que el sistema cumple con el objetivo de generar alertas de riesgo a tiempo. Adicionalmente, el Informe Final Académico documenta todo el proceso de ingeniería, sirviendo como evidencia del cumplimiento de las metas del proyecto y de la investigación."),
    ("Cierre Administrativo y Sustentación Final", "Resultados: Se concluyó formalmente el proceso de ayudantía, entregando todos los componentes técnicos (código fuente, manuales) a las partes interesadas. La sustentación garantiza la transferencia de conocimiento, permitiendo que la institución pueda continuar, mantener o escalar el prototipo en futuras fases.")
]

for title, desc in resultados:
    add_heading(doc, title, 3)
    doc.add_paragraph(desc)

add_heading(doc, 'Proyecciones y Contribución para la Continuidad de la Investigación (UNACH-LA)', 1)

doc.add_paragraph('El presente informe documenta el trabajo realizado durante nuestra participación como estudiantes en el programa de ayudantías de cátedra e investigación de este semestre. Nuestro principal objetivo fue sentar las bases tecnológicas y metodológicas para el proyecto UNACH-LA (Learning Analytics).')

doc.add_paragraph('Aunque nuestra participación corresponde a una ayudantía semestral, los resultados obtenidos aquí no son un punto final, sino el cimiento para que la Universidad Nacional de Chimborazo (UNACH) pueda dar continuidad y escalar esta investigación en el futuro. Dejamos como aporte a la institución los siguientes pilares de continuidad:')

doc.add_paragraph('1. Arquitectura y Modelos Base: Entregamos un pipeline de Machine Learning documentado y funcional, así como un prototipo de interfaz. Los futuros investigadores y ayudantes podrán tomar este repositorio, refactorizar el código, incorporar nuevas variables predictivas y conectar bases de datos más grandes sin tener que empezar desde cero.', style='List Bullet')
doc.add_paragraph('2. Lineamientos Éticos y Normativos: El análisis jurídico y los protocolos de anonimización desarrollados garantizan que cualquier investigador que retome el proyecto tenga un marco de trabajo que cumple con la Ley de Protección de Datos Personales, evitando contratiempos burocráticos.', style='List Bullet')
doc.add_paragraph('3. Escalabilidad Institucional: Al estructurar el proyecto mediante componentes independientes (Backend en Python y Frontend en React), hemos facilitado la futura integración del sistema UNACH-LA directamente con el SICO o Moodle Institucional por parte de la Dirección de Tecnologías de la Información (DTI).', style='List Bullet')

doc.add_paragraph('Confiamos en que la semilla plantada en esta ayudantía de investigación impulsará la adopción permanente de herramientas de Inteligencia Artificial para combatir la deserción estudiantil y mejorar la calidad académica en la UNACH.')

doc.save('INFORME_FINAL_COMPLETO_UNACH_LA_Actualizado.docx')
print("Documento modificado y guardado exitosamente como INFORME_FINAL_COMPLETO_UNACH_LA_Actualizado.docx")
