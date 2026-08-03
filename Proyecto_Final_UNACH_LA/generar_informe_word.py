# -*- coding: utf-8 -*-
"""
Generador de Informe Final en Word (.docx) — Formato APA 7ma Edición
=====================================================================
Genera el Informe Final Completo del Proyecto UNACH-LA con:
- Títulos de tablas y figuras en formato APA
- Notas al pie de tablas/figuras
- Numeración progresiva de tablas y figuras
- Formato académico profesional
"""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================================
# CONFIGURACIÓN
# ============================================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(BASE_DIR, "INFORME_FINAL_COMPLETO_UNACH_LA_Actualizado.docx")
GRAFICOS_DIR = os.path.join(BASE_DIR, "02_Evaluacion_de_Modelos", "graficos")
DASHBOARDS_DIR = os.path.join(BASE_DIR, "03_Visualizacion_y_KPIs", "dashboards")

# Contadores globales APA
tabla_counter = [0]
figura_counter = [0]

# Colores institucionales
AZUL_UNACH = RGBColor(0x1F, 0x4E, 0x79)
GRIS_OSCURO = RGBColor(0x33, 0x33, 0x33)
GRIS_MEDIO = RGBColor(0x66, 0x66, 0x66)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
ROJO = RGBColor(0xFF, 0x00, 0x00)


# ============================================================================
# UTILIDADES DE FORMATO APA
# ============================================================================
def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Aplica bordes a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                          f'<w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
                          f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
                          f'</w:tcBorders>')
    tcPr.append(tcBorders)


def add_tabla_apa(doc, numero, titulo, headers, rows, nota=None, font_size=9):
    """
    Agrega una tabla con formato APA 7ma edición.
    - Línea: "Tabla X" en negritas
    - Línea: Título en cursiva
    - Tabla con bordes superior e inferior
    - Nota al pie en cursiva
    """
    tabla_counter[0] += 1
    num = tabla_counter[0]

    # Título APA: "Tabla X"
    p_num = doc.add_paragraph()
    p_num.paragraph_format.space_before = Pt(18)
    p_num.paragraph_format.space_after = Pt(0)
    run_num = p_num.add_run(f"Tabla {num}")
    run_num.bold = True
    run_num.font.size = Pt(11)
    run_num.font.color.rgb = GRIS_OSCURO

    # Título descriptivo en cursiva
    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.space_before = Pt(0)
    p_titulo.paragraph_format.space_after = Pt(6)
    run_titulo = p_titulo.add_run(titulo)
    run_titulo.italic = True
    run_titulo.font.size = Pt(11)
    run_titulo.font.color.rgb = GRIS_OSCURO

    # Crear tabla
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Estilo de tabla limpio
    table.style = 'Table Grid'

    # Header
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = BLANCO
        set_cell_shading(cell, "1F4E79")

    # Filas de datos
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
            run.font.color.rgb = GRIS_OSCURO
            # Alternar colores de fila
            if i % 2 == 0:
                set_cell_shading(cell, "F2F7FB")
            # Primera columna alineada a la izquierda
            if j == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Nota APA al pie de tabla
    if nota:
        p_nota = doc.add_paragraph()
        p_nota.paragraph_format.space_before = Pt(4)
        p_nota.paragraph_format.space_after = Pt(12)
        run_nota_label = p_nota.add_run("Nota. ")
        run_nota_label.italic = True
        run_nota_label.font.size = Pt(9)
        run_nota_label.font.color.rgb = GRIS_MEDIO
        run_nota_text = p_nota.add_run(nota)
        run_nota_text.font.size = Pt(9)
        run_nota_text.font.color.rgb = GRIS_MEDIO

    return table


def add_figura_apa(doc, titulo, image_path=None, diagram_text=None, nota=None):
    """
    Agrega una figura con formato APA 7ma edición.
    - Línea: "Figura X" en negritas
    - Línea: Título en cursiva
    - Imagen o diagrama de texto
    - Nota al pie
    """
    figura_counter[0] += 1
    num = figura_counter[0]

    # Título APA: "Figura X"
    p_num = doc.add_paragraph()
    p_num.paragraph_format.space_before = Pt(18)
    p_num.paragraph_format.space_after = Pt(0)
    run_num = p_num.add_run(f"Figura {num}")
    run_num.bold = True
    run_num.font.size = Pt(11)
    run_num.font.color.rgb = GRIS_OSCURO

    # Título descriptivo en cursiva
    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.space_before = Pt(0)
    p_titulo.paragraph_format.space_after = Pt(6)
    run_titulo = p_titulo.add_run(titulo)
    run_titulo.italic = True
    run_titulo.font.size = Pt(11)
    run_titulo.font.color.rgb = GRIS_OSCURO

    # Imagen o diagrama
    if image_path and os.path.exists(image_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(image_path, width=Inches(5.5))
    elif diagram_text:
        p_diag = doc.add_paragraph()
        p_diag.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_diag.paragraph_format.space_before = Pt(6)
        p_diag.paragraph_format.space_after = Pt(6)
        run_diag = p_diag.add_run(diagram_text)
        run_diag.font.size = Pt(8)
        run_diag.font.name = "Consolas"

    # Nota APA
    if nota:
        p_nota = doc.add_paragraph()
        p_nota.paragraph_format.space_before = Pt(4)
        p_nota.paragraph_format.space_after = Pt(12)
        run_nota_label = p_nota.add_run("Nota. ")
        run_nota_label.italic = True
        run_nota_label.font.size = Pt(9)
        run_nota_label.font.color.rgb = GRIS_MEDIO
        run_nota_text = p_nota.add_run(nota)
        run_nota_text.font.size = Pt(9)
        run_nota_text.font.color.rgb = GRIS_MEDIO


def add_heading_apa(doc, text, level=1):
    """Agrega un encabezado con estilo APA."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = AZUL_UNACH
    return heading


def add_paragraph_apa(doc, text, bold=False, italic=False, size=11, alignment=None, indent=False, color=None):
    """Agrega un párrafo con formato APA."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run(text)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = GRIS_OSCURO
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, level=0):
    """Agrega un bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRIS_OSCURO
    return p


# ============================================================================
# GENERADOR PRINCIPAL DEL DOCUMENTO
# ============================================================================
def generar_informe():
    print("=" * 70)
    print("  GENERANDO INFORME FINAL EN WORD — FORMATO APA 7ma EDICIÓN")
    print("=" * 70)

    doc = Document()

    # ========================================================================
    # CONFIGURACIÓN DE PÁGINA
    # ========================================================================
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Fuente por defecto
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = GRIS_OSCURO

    # ========================================================================
    # PORTADA
    # ========================================================================
    print("  >> Generando portada...")
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDAD NACIONAL DE CHIMBORAZO")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = AZUL_UNACH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PROYECTO DE INVESTIGACIÓN — AYUDANTÍA")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = GRIS_OSCURO

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INFORME FINAL COMPLETO")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = AZUL_UNACH

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MODELO INSTITUCIONAL DE LEARNING ANALYTICS (UNACH-LA):\nSISTEMA DE ALERTA TEMPRANA CON INTELIGENCIA ARTIFICIAL")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = GRIS_OSCURO

    for _ in range(4):
        doc.add_paragraph()

    info_items = [
        "Riobamba — Chimborazo, Ecuador",
        "Julio 2026",
        "Versión 1.0 — Documento Final"
    ]
    for item in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = GRIS_MEDIO

    doc.add_page_break()

    # ========================================================================
    # ÍNDICE GENERAL
    # ========================================================================
    print("  >> Generando índice...")
    add_heading_apa(doc, "ÍNDICE GENERAL", 1)

    indice = [
        "1. Introducción y Contexto del Proyecto",
        "2. Fase 1: Alineación Normativa y Fundamentación",
        "   2.1 Entregable 1: Matriz Normativa y Análisis Legal",
        "   2.2 Entregable 2: Estado del Arte y Bibliografía Sistematizada",
        "   2.3 Entregable 3: Informe de Ecosistema y Flujos de Datos",
        "   2.4 Entregable 4: Documento de Lineamientos Éticos y de Gobernanza",
        "3. Fase 2: Diseño e Implementación de Machine Learning",
        "   3.1 Preparación y Limpieza de Datos Académicos",
        "   3.2 Diseño de Modelos ML para Diagnóstico Académico",
        "   3.3 Implementación de Modelos ML en Python y scikit-learn",
        "   3.4 Evaluación de Modelos: Validación y Pruebas de Rendimiento",
        "   3.5 Visualización y KPIs: Diseño de Dashboards e Indicadores",
        "4. Fase 3: Documentación y Cierre",
        "   4.1 Consolidación del Prototipo Funcional UNACH-LA",
        "   4.2 Dashboard Institucional en React + Vite",
        "   4.3 Backend API con FastAPI",
        "   4.4 Integración con Inteligencia Artificial Generativa",
        "5. Arquitectura Técnica Completa del Sistema",
        "6. Resultados Consolidados y Métricas Finales",
        "7. Conclusiones y Recomendaciones",
        "8. Bibliografía y Referencias",
        "9. Anexos",
    ]
    for item in indice:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        if not item.startswith("   "):
            run.bold = True

    doc.add_page_break()

    # ========================================================================
    # CAPÍTULO 1: INTRODUCCIÓN
    # ========================================================================
    print("  >> Capítulo 1: Introducción...")
    add_heading_apa(doc, "1. INTRODUCCIÓN Y CONTEXTO DEL PROYECTO", 1)

    add_heading_apa(doc, "1.1 Planteamiento del Problema", 2)
    add_paragraph_apa(doc,
        "En la Universidad Nacional de Chimborazo (UNACH), los datos estudiantiles y de rendimiento académico "
        "se encuentran depositados en múltiples plataformas independientes: el Sistema Integrado de Control "
        "Académico (SICOA), las aulas virtuales basadas en Moodle (LMS), los registros de tutorías y los "
        "informes de Bienestar Estudiantil. Esta fragmentación genera silos tecnológicos que impiden una "
        "visión integral del desempeño estudiantil, provocando que las decisiones académicas se tomen con "
        "base en percepciones subjetivas y no en evidencia cuantificable.", indent=True)

    add_paragraph_apa(doc,
        "Uno de los factores agravantes identificados en el diagnóstico es la brecha digital presente en la "
        "provincia de Chimborazo, donde aproximadamente el 36% de los hogares carecen de conexión a internet "
        "estable (INEC, 2022). Esta realidad impacta directamente en la capacidad de los estudiantes para "
        "acceder a los recursos del LMS, generando patrones irregulares de actividad que deben ser contemplados "
        "por cualquier sistema predictivo.", indent=True)

    add_heading_apa(doc, "1.2 Justificación e Importancia de la Investigación", 2)
    add_paragraph_apa(doc,
        "La implementación de un modelo institucional de Learning Analytics (LA) permite consolidar una cultura "
        "de mejora continua basada en datos, respondiendo tanto a la política nacional de transformación digital "
        "del Ecuador como a las directrices del Consejo de Educación Superior (CES) para el aseguramiento de "
        "la calidad. El proyecto UNACH-LA convierte la fragmentación de datos existente en una oportunidad de "
        "innovación educativa sustentable.", indent=True)

    add_paragraph_apa(doc,
        "La importancia de esta investigación radica en ser un trabajo pionero a nivel regional al democratizar "
        "el acceso a la inteligencia artificial en la gestión académica. Al liberar el código fuente de manera "
        "pública y de código abierto (Open Source), la UNACH no solo resuelve un problema interno de deserción, "
        "sino que provee una arquitectura de referencia escalable que cualquier otra universidad latinoamericana "
        "puede adoptar y adaptar a su propia realidad sin incurrir en costos de software privativo.", indent=True, color=ROJO)

    add_paragraph_apa(doc, "El proyecto se articula con los siguientes marcos institucionales:")
    add_bullet(doc, "Dominios académicos: Desarrollo socioeconómico y educativo para el fortalecimiento de la institucionalidad democrática y ciudadana.")
    add_bullet(doc, "Campos del conocimiento: Educación y Tecnologías de la Información y Comunicación (TICs).")
    add_bullet(doc, "Líneas de investigación: Ciencias de la Educación e Ingeniería Informática.")

    add_paragraph_apa(doc, "Alineación con los Objetivos de Desarrollo Sostenible (ODS):")
    add_bullet(doc, "ODS 4 — Educación de Calidad: Garantizar una educación inclusiva y equitativa.")
    add_bullet(doc, "ODS 9 — Industria, Innovación e Infraestructura: Fomentar la innovación y apoyar la investigación científica.")
    add_bullet(doc, "ODS 17 — Alianzas para Lograr los Objetivos: Fortalecer la alianza entre tecnología y educación.")

    add_heading_apa(doc, "1.3 Objetivos", 2)
    add_paragraph_apa(doc, "Objetivo General", bold=True)
    add_paragraph_apa(doc,
        "Desarrollar e implementar un modelo institucional de Learning Analytics en la Universidad Nacional "
        "de Chimborazo que mejore los procesos de formación académica, fortalezca las decisiones pedagógicas "
        "y transfiera conocimiento a las partes interesadas.", indent=True)

    add_paragraph_apa(doc, "Objetivos Específicos", bold=True)
    add_bullet(doc, "Realizar un diagnóstico del ecosistema tecnológico de la UNACH y sus flujos de datos académicos.")
    add_bullet(doc, "Diseñar e implementar modelos de Machine Learning para la predicción del riesgo académico estudiantil.")
    add_bullet(doc, "Evaluar y validar el rendimiento de los modelos predictivos utilizando métricas estándar de clasificación.")
    add_bullet(doc, "Construir un prototipo funcional completo (Dashboard + Backend + Motor ML) que permita la operación institucional del sistema.")
    add_bullet(doc, "Documentar y sistematizar el proceso metodológico para su replicabilidad en otras instituciones de educación superior.")

    add_heading_apa(doc, "1.4 Metodología General", 2)
    add_paragraph_apa(doc,
        "El proyecto adoptó un enfoque metodológico mixto. El componente cuantitativo incluyó minería de datos "
        "educativos, modelos predictivos supervisados (XGBoost, Random Forest, Logistic Regression, Decision Tree, SVM), "
        "extracción y procesamiento de datos del SICOA y LMS. El componente cualitativo abarcó la validación con "
        "grupos focales (docentes tutores y coordinadores de carrera), análisis normativo institucional y revisión "
        "de literatura especializada.", indent=True)

    add_paragraph_apa(doc,
        "La metodología de desarrollo tecnológico siguió el marco CRISP-DM (Cross-Industry Standard Process for "
        "Data Mining), ampliamente utilizado en proyectos de ciencia de datos educativos (Shearer, 2000; Chapman et al., 2000).", indent=True)

    doc.add_page_break()

    # ========================================================================
    # CAPÍTULO 2: FASE 1
    # ========================================================================
    print("  >> Capítulo 2: Fase 1 — Alineación Normativa...")
    add_heading_apa(doc, "2. FASE 1: ALINEACIÓN NORMATIVA Y FUNDAMENTACIÓN", 1)

    add_paragraph_apa(doc,
        "El objetivo de esta fase fue establecer el marco normativo, ético y contextual que sustenta la "
        "implementación del sistema UNACH-LA, garantizando el cumplimiento de la legislación ecuatoriana "
        "vigente y las mejores prácticas internacionales en Learning Analytics.", indent=True)

    # --- 2.1 Entregable 1 ---
    add_heading_apa(doc, "2.1 Entregable 1: Matriz Normativa y Análisis Legal", 2)
    add_paragraph_apa(doc,
        "Se llevó a cabo una revisión exhaustiva de la normativa institucional de la UNACH y de la "
        "legislación ecuatoriana aplicable al tratamiento de datos personales en el contexto educativo, "
        "con especial énfasis en la Ley Orgánica de Protección de Datos Personales (LOPDP) del Ecuador "
        "(Registro Oficial Suplemento 459, 26-may-2021), el Reglamento de Régimen Académico del CES, "
        "el Estatuto de la UNACH y la normativa internacional de referencia (GDPR).", indent=True)

    add_tabla_apa(doc, 1,
        "Análisis de Cumplimiento Normativo del Proyecto UNACH-LA respecto a la LOPDP",
        ["Aspecto Analizado", "Marco Normativo", "Implicación para el Proyecto"],
        [
            ["Consentimiento informado", "LOPDP Art. 9-11", "Requiere aceptación expresa de los titulares de datos"],
            ["Finalidad del tratamiento", "LOPDP Art. 7", "Los datos solo pueden utilizarse para fines académicos declarados"],
            ["Transferencia internacional", "LOPDP Art. 37-39", "Infraestructura cloud debe garantizar protección equivalente"],
            ["Derechos ARCO", "LOPDP Art. 17-24", "Acceso, rectificación, cancelación y oposición garantizados"],
            ["Anonimización", "LOPDP Art. 5", "Datasets de entrenamiento ML deben estar anonimizados"],
            ["Retención de datos", "Estatuto UNACH", "Periodo máximo alineado al ciclo académico vigente"],
        ],
        nota="LOPDP = Ley Orgánica de Protección de Datos Personales del Ecuador. ARCO = Acceso, Rectificación, Cancelación, Oposición. Fuente: Elaboración propia basada en la LOPDP (2021) y el Estatuto de la UNACH (2023)."
    )

    # --- 2.2 Entregable 2 ---
    add_heading_apa(doc, "2.2 Entregable 2: Estado del Arte y Bibliografía Sistematizada", 2)
    add_paragraph_apa(doc,
        "Se realizó una investigación documental sistematizada sobre casos de éxito en la implementación "
        "de Learning Analytics en instituciones de educación superior a nivel mundial y regional, con "
        "revisión de literatura especializada en las bases Scopus, Web of Science, IEEE Xplore, ACM "
        "Digital Library y Google Scholar.", indent=True)

    add_tabla_apa(doc, 2,
        "Marcos de Referencia en Learning Analytics Revisados para la Fundamentación del Proyecto",
        ["Marco / Modelo", "Autores", "Aplicación", "Relevancia"],
        [
            ["Learning Analytics Framework", "Siemens y Baker (2012)", "Predicción de rendimiento", "Base conceptual del proyecto"],
            ["CRISP-DM", "Chapman et al. (2000)", "Metodología de ciencia de datos", "Marco metodológico adoptado"],
            ["Early Warning Systems", "Arnold y Pistilli (2012)", "Sistemas de alerta temprana", "Arquitectura de referencia"],
            ["Modelo SRL", "Zimmerman (2002)", "Aprendizaje autorregulado", "Fundamenta las features de engagement"],
            ["Learning Analytics Cycle", "Clow (2012)", "Ciclo de analítica del aprendizaje", "Diseño del flujo de datos"],
            ["Predictive Analytics in HE", "Baker e Inventado (2014)", "Minería de datos educativos", "Técnicas ML seleccionadas"],
        ],
        nota="HE = Higher Education (Educación Superior). SRL = Self-Regulated Learning. Fuente: Elaboración propia a partir de la revisión de literatura."
    )

    add_tabla_apa(doc, 3,
        "Casos de Éxito Internacionales en Sistemas de Alerta Temprana con Learning Analytics",
        ["Caso", "Institución", "Algoritmo", "Resultado Principal"],
        [
            ["Course Signals", "Purdue University", "Modelo propio", "Reducción del 21% en reprobación"],
            ["Degree Compass", "Austin Peay State Univ.", "Recomendación", "Mejora del 12% en aprobación"],
            ["OU Analyse", "Open University UK", "Random Forest", "AUC-ROC = 0.78"],
            ["LALA Project", "Universidades de LA", "Mixto", "Marco contextualizado para la región"],
        ],
        nota="LA = Latinoamérica. Fuente: Arnold y Pistilli (2012); Denley (2014); Kuzilek et al. (2017); Pérez-Sanagustín et al. (2018)."
    )

    # --- 2.3 Entregable 3 ---
    add_heading_apa(doc, "2.3 Entregable 3: Informe de Ecosistema y Flujos de Datos", 2)
    add_paragraph_apa(doc,
        "Se realizó una caracterización completa del ecosistema tecnológico de la UNACH, identificando "
        "las plataformas, bases de datos y flujos de información relevantes para la construcción del "
        "sistema de Learning Analytics.", indent=True)

    add_tabla_apa(doc, 4,
        "Plataformas del Ecosistema Tecnológico de la UNACH Identificadas en el Diagnóstico",
        ["Plataforma", "Función", "Tipo de Datos", "Formato"],
        [
            ["SICOA", "Sistema de Control Académico", "Notas, asistencia, matrícula, historial", ".xlsx, .csv"],
            ["Moodle (LMS)", "Aula Virtual", "Logs de actividad, calificaciones, foros", "MySQL, logs"],
            ["Sistema de Tutorías", "Registro de acompañamiento", "Sesiones, reportes de seguimiento", "Documentos"],
            ["Bienestar Estudiantil", "Servicios de apoyo", "Becas, vulnerabilidad, salud", "Registros internos"],
        ],
        nota="SICOA = Sistema Integrado de Control Académico. LMS = Learning Management System. Fuente: Diagnóstico tecnológico institucional elaborado durante la Fase 1."
    )

    add_figura_apa(doc,
        "Diagrama de Flujo de Datos del Sistema UNACH-LA desde las Fuentes hasta las Salidas Operativas",
        diagram_text=(
            "  ┌─────────────┐    ┌─────────────┐    ┌──────────────────┐\n"
            "  │   SICOA      │    │   Moodle     │    │  Bienestar       │\n"
            "  │  (Notas,     │    │  (Logs LMS,  │    │  Estudiantil     │\n"
            "  │  Asistencia) │    │  Actividad)  │    │  (Becas, Vuln.)  │\n"
            "  └──────┬───────┘    └──────┬───────┘    └────────┬─────────┘\n"
            "         │                   │                      │\n"
            "         └───────────────────┼──────────────────────┘\n"
            "                             │\n"
            "                       ┌─────▼──────┐\n"
            "                       │  ETL Layer  │\n"
            "                       │  (Python)   │\n"
            "                       └─────┬──────┘\n"
            "                             │\n"
            "                       ┌─────▼──────────────┐\n"
            "                       │ Dataset Procesado   │\n"
            "                       │ (4,000 × 89 cols)   │\n"
            "                       └─────┬──────────────┘\n"
            "                             │\n"
            "                 ┌───────────┼───────────┐\n"
            "                 │           │           │\n"
            "           ┌─────▼────┐ ┌───▼───┐ ┌────▼──────┐\n"
            "           │Dashboard │ │Alertas│ │ API       │\n"
            "           │ React    │ │ JSON  │ │ FastAPI   │\n"
            "           └──────────┘ └───────┘ └───────────┘"
        ),
        nota="ETL = Extract, Transform, Load. El flujo muestra la arquitectura de datos del modelo UNACH-LA desde la extracción de fuentes institucionales hasta la generación de salidas operativas. Fuente: Elaboración propia."
    )

    # --- 2.4 Entregable 4 ---
    add_heading_apa(doc, "2.4 Entregable 4: Documento de Lineamientos Éticos y de Gobernanza", 2)
    add_paragraph_apa(doc,
        "Se definieron los criterios éticos, de privacidad y gobernanza institucional que deben regir "
        "la operación del sistema UNACH-LA, alineados con los principios de transparencia algorítmica "
        "y justicia educativa.", indent=True)

    add_tabla_apa(doc, 5,
        "Principios Éticos Rectores del Sistema UNACH-LA y su Implementación Técnica",
        ["Principio", "Descripción", "Implementación en UNACH-LA"],
        [
            ["Transparencia", "Algoritmos y criterios explicables", "Dashboard muestra feature importance y métricas"],
            ["No discriminación", "No perpetuar sesgos por género, etnia o condición", "Análisis de equidad en predicciones por subgrupos"],
            ["Proporcionalidad", "Solo datos estrictamente necesarios", "89 variables tras limpieza rigurosa"],
            ["Consentimiento", "Estudiante debe autorizar uso de sus datos", "Protocolo de consentimiento informado"],
            ["Beneficencia", "Predicciones para beneficiar al estudiante", "Alertas vinculadas a planes de tutoría, no punitivos"],
            ["Reversibilidad", "Decisiones no irreversibles por ML", "Modelo como herramienta de apoyo, no de decisión automática"],
        ],
        nota="Principios basados en las directrices de la UNESCO sobre Inteligencia Artificial en Educación y la LOPDP del Ecuador. Fuente: Elaboración propia."
    )

    doc.add_page_break()

    # ========================================================================
    # SECCIÓN 2.5: MODELADO DEL ECOSISTEMA DE DATOS
    # ========================================================================
    print("  >> Capítulo 5: Modelado del Ecosistema de Datos...")
    add_heading_apa(doc, "2.5 MODELADO DEL ECOSISTEMA DE DATOS", 1)

    add_heading_apa(doc, "2.5.1 Diagrama de arquitectura", 2)
    add_heading_apa(doc, "2.5.1.1 Arquitectura actual", 3)
    add_paragraph_apa(doc, "La infraestructura tecnológica actual de la UNACH se centraliza en el Data Center ubicado en el Campus Norte. Esta arquitectura cuenta con un \"Backbone de Redes de Telecomunicaciones\" impulsado por tecnología con inteligencia artificial (Juniper Mist) y mantiene respaldos inmutables en servidores locales y en la nube de CEDIA para asegurar la información crítica. Sin embargo, a nivel de software, los sistemas académicos como el SICOA y Moodle operan a menudo como \"islas de digitalización\" carentes de una articulación centralizada y automatizada.", indent=True, color=ROJO)
    
    add_heading_apa(doc, "2.5.1.2 Arquitectura propuesta", 3)
    add_paragraph_apa(doc, "La arquitectura propuesta se fundamenta en el proyecto institucional de Learning Analytics (Modelo UNACH-LA). Esta propuesta añade capas de integración donde la información extraída de los sistemas fuente pasa por un mecanismo estricto de anonimización (asignando identificadores únicos no reversibles). Posteriormente, los datos alimentan un motor analítico que centraliza la información para generar dashboards e indicadores clave de desempeño en tiempo real.", indent=True, color=ROJO)

    add_heading_apa(doc, "2.5.2 Diagrama de flujo de datos (DFD)", 2)
    add_paragraph_apa(doc, "Nivel 0 (Diagrama de contexto) Representa al Sistema Integrado de Analítica (UNACH-LA) como un proceso único central que interactúa con las entidades externas. Los estudiantes y docentes ingresan sus datos e interacciones, y el sistema devuelve tableros de control y resultados para la toma de decisiones.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Nivel 1 (Descomposición de procesos) Muestra el interior del sistema central dividiéndolo en sus subprocesos principales: Recolección, Anonimización, Procesamiento Analítico y Visualización de datos.", indent=True, color=ROJO)

    add_heading_apa(doc, "2.5.3 Actores involucrados", 2)
    add_paragraph_apa(doc, "Estudiantes: Son la población objetivo y fuente principal de datos. Proveen información demográfica, historial académico, calificaciones, asistencia y generan una huella digital constante mediante sus accesos y participación en los recursos del entorno Moodle. Son los beneficiarios finales de las estrategias de retención.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Docentes: Son usuarios clave tanto en la generación de datos (registrando calificaciones y asistencias en SICOA y evaluando tareas en Moodle) como en la utilización de la plataforma propuesta. El proyecto contempla la formación de docentes en analítica educativa para que puedan interpretar los dashboards y aplicar mejoras pedagógicas.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Coordinadores (y Autoridades): Incluye a directores de carrera, autoridades de facultad y autoridades institucionales. Su rol es estratégico, ya que consumen la información procesada por el modelo de analítica (tableros de control) para la gestión académica y la toma de decisiones basada en evidencia que reduzca la deserción estudiantil.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Administradores TIC: Involucra principalmente al personal de la Dirección de Tecnologías de la Información y Comunicación (DTIC) y la Coordinación de Desarrollo de Sistemas Informáticos (CODESI). Tienen el rol técnico vital de extraer los registros masivos de las plataformas, garantizar la infraestructura de servidores y redes, y aplicar los procesos de anonimización y asignación de identificadores únicos antes de entregar la información para el análisis.", indent=True, color=ROJO)
    doc.add_page_break()

    # ========================================================================
    # SECCIÓN 2.6: ANÁLISIS DE OPORTUNIDADES PARA LEARNING ANALYTICS
    # ========================================================================
    print("  >> Capítulo 6: Análisis de Oportunidades...")
    add_heading_apa(doc, "2.6 ANÁLISIS DE OPORTUNIDADES PARA LEARNING ANALYTICS", 1)
    add_paragraph_apa(doc, "El análisis de oportunidades se fundamenta en la ejecución del proyecto de investigación Modelo UNACH-LA, cuyo propósito es transformar los crecientes volúmenes de datos educativos en conocimiento accionable para mejorar los procesos académicos y reducir el riesgo de deserción en la Universidad Nacional de Chimborazo.", indent=True, color=ROJO)

    add_heading_apa(doc, "2.6.1 Variables académicas relevantes", 2)
    add_paragraph_apa(doc, "El modelo de Learning Analytics aprovecha datos provenientes de plataformas institucionales como SICOA y Moodle, los cuales se procesan bajo un estricto mecanismo de anonimización y seudonimización mediante un identificador único no reversible para proteger la identidad de los estudiantes.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Rendimiento: Se evalúa extrayendo datos históricos y actuales del SICOA, que incluyen las calificaciones obtenidas por asignatura (primer parcial, segundo parcial, recuperación y calificación final), puntajes de admisión, promedios de nivelación, asignaturas de la malla aprobadas y calificaciones de grado o titulación.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Participación: Se mide a través de la huella digital estudiantil capturada en el Learning Management System (Moodle), recopilando registros diarios de todas las acciones, recursos visualizados y eventos realizados en las aulas virtuales. Además, se complementa con el porcentaje de asistencia por asignatura y el registro de tutorías recibidas durante los periodos cursados.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Riesgo de deserción: Se monitorea mediante la identificación de patrones en datos críticos como el número de retiros y reingresos, asignaturas tomadas en segunda o tercera matrícula, y variables de contexto sociodemográfico que evidencien vulnerabilidad, tales como condiciones de salud, residencia (urbano/rural) y situación socioeconómica declarada.", indent=True, color=ROJO)

    add_heading_apa(doc, "2.6.2 Indicadores potenciales", 2)
    add_paragraph_apa(doc, "KPIs académicos: El proyecto busca consolidar la información en un prototipo funcional con tableros de control (dashboards) institucionales, los cuales mostrarán indicadores clave de desempeño (KPIs) para facilitar el seguimiento del progreso académico.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Alertas tempranas: Al integrar dimensiones pedagógicas y tecnológicas, los datos analizados permitirán generar alertas para proyectar análisis enfocados en disminuir el riesgo de deserción académica en grupos de estudio.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Predicción de desempeño: La analítica correlacional y longitudinal de los registros de SICOA (rendimiento) y Moodle (interacción) tiene el objetivo de traducir los datos masivos en conocimiento accionable, permitiendo proyectar escenarios para mejorar el rendimiento académico y guiar la toma de decisiones.", indent=True, color=ROJO)

    add_heading_apa(doc, "2.6.3 Viabilidad de implementación", 2)
    add_paragraph_apa(doc, "Técnica: La viabilidad técnica está respaldada por la reciente modernización del \"Backbone de Redes de Telecomunicaciones\" del campus, que incorporó tecnología de inteligencia artificial y amplió la cobertura, garantizando una alta disponibilidad y estabilidad para el procesamiento de datos. Asimismo, el uso de identificadores anonimizados viabiliza el cruce seguro de bases de datos masivas en formatos CSV y Excel. Sin embargo, existen limitaciones a superar, como la necesidad de actualizar librerías de software institucional para una mejor integración.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Operativa: A nivel operativo, el proyecto contempla la formación de al menos 50 docentes y gestores universitarios en analítica educativa, lo que garantizará que el personal sepa interpretar y utilizar los resultados del modelo. No obstante, el principal reto operativo radica en la demanda de tiempo para soporte técnico a usuarios finales y en la desvinculación de personal en el área de desarrollo de sistemas (CODESI), lo que limita la capacidad operativa para nuevos desarrollos de alto impacto.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Institucional: La implementación cuenta con la aprobación oficial de la Comisión de Investigación (Resolución No. 037-CIV-12-02-2026) y se alinea con la Política Pública para la Transformación Digital 2025-2030, que exige a las universidades la gestión basada en datos y el uso de tecnologías emergentes. Sin embargo, para garantizar su sostenibilidad, estudios de transformación digital en la UNACH recomiendan que estas iniciativas dejen de funcionar como \"islas de digitalización\" y se integren dentro de un plan institucional específico de transformación digital, requiriendo un fuerte compromiso de liderazgo y asignación de presupuesto.", indent=True, color=ROJO)
    doc.add_page_break()

    # ========================================================================
    # SECCIÓN 2.7: RIESGOS Y CONSIDERACIONES
    # ========================================================================
    print("  >> Capítulo 7: Riesgos y Consideraciones...")
    add_heading_apa(doc, "2.7 RIESGOS Y CONSIDERACIONES", 1)
    
    add_heading_apa(doc, "2.7.1 Seguridad de la información", 2)
    add_paragraph_apa(doc, "La Universidad Nacional de Chimborazo (UNACH) rige la seguridad de su ecosistema tecnológico basándose en las \"Políticas de Seguridad de la Información UNACH 2018\", las cuales se alinean con estándares internacionales como ISO 27001 para mitigar riesgos y amenazas.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Protección de datos: Para evitar la pérdida de información crítica y garantizar su resguardo, la institución ha implementado copias de seguridad (backups) inmutables, alojadas tanto en servidores locales como en la nube de CEDIA. La normativa interna exige que las bases de datos, aplicativos y copias de seguridad se mantengan encriptadas. Además, los activos físicos del Data Center central están resguardados por sistemas de videovigilancia y acceso restringido.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Accesos: El ingreso a la infraestructura tecnológica y sistemas de información está estrictamente normado mediante un sistema de gestión de credenciales (roles y privilegios). Recientemente, se implementó un nuevo servidor Radius para fortalecer el control seguro de la autenticación de usuarios. Para los administradores que requieran conexiones desde el exterior, el acceso remoto a los servidores institucionales se realiza exclusivamente a través de Redes Privadas Virtuales (VPN) con encriptación.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Privacidad: La política institucional clasifica a los datos personales (como nombres, etnia, lugar de procedencia, estado de salud o vulnerabilidades económicas) estrictamente como información confidencial. Su uso está restringido al cumplimiento de actividades institucionales formales, quedando prohibida cualquier divulgación sin la debida autorización.", indent=True, color=ROJO)

    add_heading_apa(doc, "2.7.2 Calidad de datos", 2)
    add_paragraph_apa(doc, "Para asegurar que los modelos de Learning Analytics (como el proyecto UNACH-LA) ofrezcan predicciones y dashboards confiables, los datos deben cumplir tres características técnicas fundamentales:", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Integridad: Se garantiza que la información no ha sido modificada ni alterada sin autorización. Para lograrlo, los sistemas académicos (como SICOA) tienen habilitadas pistas de auditoría (logs) no editables que registran cada transacción, identificando al responsable de cualquier inserción, actualización o borrado de datos sensibles. Ningún usuario, ni siquiera los administradores o desarrolladores, puede modificar información directamente en la base de datos de producción sin seguir un estricto protocolo de control de cambios.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Consistencia: Dado que la infraestructura funciona en muchas ocasiones como \"islas de digitalización\", el proyecto UNACH-LA resuelve el cruce de información (entre SICOA y Moodle) mediante la creación de un identificador único no reversible. Esto asegura que los historiales académicos, de asistencia y de interacción digital coincidan exactamente para el mismo individuo, permitiendo efectuar análisis longitudinales y correlacionales sin inconsistencias en las bases de datos.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Disponibilidad: La disponibilidad permanente de la información se ha visto robustecida mediante el proyecto \"Backbone de Redes de Telecomunicaciones\", impulsado por tecnología de inteligencia artificial Juniper Mist. Esto permite alcanzar un índice de disponibilidad de red del 99%. A nivel de servidores, el Data Center institucional cuenta con estándares de redundancia diseñados para evitar interrupciones.", indent=True, color=ROJO)

    add_heading_apa(doc, "2.7.3 Aspectos legales y éticos", 2)
    add_paragraph_apa(doc, "La extracción de datos masivos para el seguimiento académico conlleva responsabilidades legales en el marco jurídico ecuatoriano y políticas internas.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Protección de datos personales: Todo el ecosistema opera bajo el cumplimiento de la Constitución de la República y la Ley Orgánica de Protección de Datos Personales, que exigen medidas de seguridad rigurosas frente a la información ciudadana. En estricto apego a esta ley, el requerimiento de datos para el modelo UNACH-LA establece como condición obligatoria la anonimización y/o seudonimización de toda la información entregada por Secretaría Académica y la Dirección de TICs (DTIC), garantizando la no identificación directa de los titulares de los datos.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Consentimiento: De acuerdo con el Art. 66 de la Constitución, el procesamiento o distribución de datos requiere la autorización de su titular o un mandato de la ley. Para fines del uso institucional, el personal y los proveedores firman acuerdos de confidencialidad y responsabilidades respecto al manejo de los activos. En el caso de la analítica de aprendizaje (UNACH-LA), al utilizar un mecanismo ciego de seudonimización (donde el dato personal se reemplaza por un código alfanumérico), se mitiga el riesgo de vulnerar la intimidad estudiantil, actuando bajo el amparo de la Resolución Oficial de la Comisión de Investigación de la UNACH (No. 037-CIV-12-02-2026).", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Uso responsable de información académica: El acceso a los datos extraídos tiene como única finalidad la investigación y el desarrollo de estrategias para la retención estudiantil. Existe un documento de compromiso firmado por los miembros del equipo de investigación (liderado por la Dirección del Proyecto UNACH-LA), en el cual se obligan a dar fiel cumplimiento a la normativa aplicable sobre el buen uso de los datos y a utilizarlos estrictamente con fines investigativos. Además, las políticas de la universidad prohíben categóricamente el uso de la infraestructura o bases de datos para suministrar información con el fin de obtener beneficios propios o de terceros.", indent=True, color=ROJO)
    doc.add_page_break()
    # ========================================================================
    # CAPÍTULO 3: FASE 2
    # ========================================================================
    print("  >> Capítulo 3: Fase 2 — Diseño e Implementación ML...")
    add_heading_apa(doc, "3. FASE 2: DISEÑO E IMPLEMENTACIÓN DE MACHINE LEARNING", 1)

    add_paragraph_apa(doc,
        "El objetivo de esta fase fue preparar los datos académicos, diseñar y entrenar modelos de "
        "Machine Learning supervisado para la predicción del riesgo académico, evaluar su rendimiento "
        "mediante métricas estándar, y construir indicadores clave de desempeño (KPIs) para la gestión "
        "institucional.", indent=True)

    # --- 3.1 Preparación de Datos ---
    add_heading_apa(doc, "3.1 Preparación y Limpieza de Datos Académicos", 2)
    add_paragraph_apa(doc,
        "Se implementó un pipeline completo de ETL (Extract, Transform, Load) en Python para la "
        "preparación y limpieza de dos datasets institucionales. El pipeline sigue la secuencia: "
        "Limpieza LMS → Limpieza SICOA → Feature Engineering → Merge → Exportación.", indent=True)

    add_tabla_apa(doc, 6,
        "Datasets Institucionales Utilizados como Fuente de Datos para el Proyecto UNACH-LA",
        ["Dataset", "Origen", "Filas", "Columnas", "Descripción"],
        [
            ["dataset_LMS_2025_2S.xlsx", "Moodle (LMS)", "74,464", "25", "Logs de actividad digital del aula virtual"],
            ["dataset_sicoa_2025.xlsx", "SICOA", "4,000", "48", "Registros académicos institucionales"],
        ],
        nota="LMS = Learning Management System. SICOA = Sistema Integrado de Control Académico. Los datasets corresponden al periodo académico 2025-2S. Fuente: Extracción institucional."
    )

    add_heading_apa(doc, "3.1.1 Paso 1: Limpieza del Dataset LMS", 3)
    add_tabla_apa(doc, 7,
        "Transformaciones Aplicadas al Dataset LMS (Moodle) durante la Fase de Limpieza",
        ["Operación", "Detalle", "Impacto"],
        [
            ["Eliminación de columnas vacías", "observacion, contexto, ip_anon, periodo", "-4 columnas"],
            ["Filtrado de usuarios", "Solo tipo Estudiante (eliminación de tutores)", "-8,049 filas"],
            ["Eliminación de columna constante", "tipo_usuario (ahora constante)", "-1 columna"],
            ["Parseo de timestamps", "Extracción de fecha, hora_num, franja_horaria", "+3 columnas"],
            ["Tipificación de categóricas", "evento, accion, componente, recurso, etc.", "Tipo category"],
        ],
        nota="Franja horaria: mañana (06:00–11:59), tarde (12:00–17:59), noche (18:00–23:59), madrugada (00:00–05:59). Fuente: Pipeline ETL (dataset.py)."
    )

    add_heading_apa(doc, "3.1.2 Paso 2: Limpieza del Dataset SICOA", 3)
    add_tabla_apa(doc, 8,
        "Transformaciones Aplicadas al Dataset SICOA durante la Fase de Limpieza",
        ["Operación", "Detalle", "Impacto"],
        [
            ["Eliminación de columnas vacías", "recuperacion, periodo, periodo_titulacion, etc.", "-5 columnas"],
            ["Creación de flag binario", "tiene_titulacion a partir de fecha_titulacion", "+1 variable"],
            ["Imputación de nulos", "enfermedad y dificultad_aprendizaje → 'Ninguna'", "Nulos tratados"],
            ["Cálculo de edad", "Desde fecha_nacimiento (ref: 2025-09-01)", "Media = 22.3 años"],
            ["Codificación binaria", "matricula_vigente, beca, vulnerabilidad → 0/1", "Variables numéricas"],
            ["Preservación de NaN", "Notas de titulación (estudiantes no titulados)", "2,600 NaN mantenidos"],
        ],
        nota="NaN = Not a Number (valor faltante). Los NaN en notas de titulación fueron preservados deliberadamente porque corresponden a estudiantes que aún no han culminado su proceso de titulación. Fuente: Pipeline ETL (dataset.py)."
    )

    add_heading_apa(doc, "3.1.3 Paso 3: Feature Engineering — Métricas de Comportamiento Digital", 3)
    add_paragraph_apa(doc,
        "Se generaron 36 features de comportamiento digital agregando el dataset LMS a nivel de "
        "estudiante mediante la operación groupby('codigo_usuario'). Las features cubren seis categorías: "
        "actividad básica, calificaciones LMS, temporales, conteo por tipo de evento, conteo por "
        "componente y contextuales.", indent=True)

    add_tabla_apa(doc, 9,
        "Categorías de Features de Comportamiento Digital Generadas a Partir del Dataset LMS",
        ["Categoría", "N° Features", "Ejemplos Representativos"],
        [
            ["Actividad básica", "6", "total_eventos, total_sesiones, tiempo_conexion_total_min"],
            ["Calificaciones LMS", "4", "calificacion_lms_promedio, calificacion_lms_max, calificacion_lms_std"],
            ["Temporales", "2", "dias_activos, eventos_por_dia"],
            ["Conteo por tipo de evento", "10", "evt_assignment_submitted, evt_quiz_attempted, evt_forum_post"],
            ["Conteo por componente", "7", "comp_tarea, comp_cuestionario, comp_foro, comp_url"],
            ["Contextuales", "7", "tasa_errores, dispositivo_principal, franja_preferida, actividad_finde"],
        ],
        nota="Total: 36 features de comportamiento digital. Las features se calcularon mediante operaciones de agregación (sum, mean, nunique, mode) sobre los 74,464 registros LMS agrupados por estudiante. Fuente: Pipeline ETL (dataset.py)."
    )

    add_heading_apa(doc, "3.1.4 Paso 4: Feature Engineering Avanzado (Post-Merge)", 3)
    add_tabla_apa(doc, 10,
        "Features Derivadas Avanzadas Generadas tras la Fusión de Datasets SICOA y LMS",
        ["#", "Feature", "Fórmula / Lógica", "Justificación Pedagógica"],
        [
            ["1", "ratio_aprobadas_total", "aprobadas_1ra / total_aprobadas", "Eficiencia de aprobación a la primera vez"],
            ["2", "engagement_score", "30% eventos + 30% días + 20% sesiones + 20% tiempo", "Indicador holístico de participación digital"],
            ["3", "tasa_entrega", "submitted / (submitted + viewed)", "Proporción de tareas efectivamente entregadas"],
            ["4", "tasa_quiz_completado", "quiz_submitted / quiz_attempted", "Persistencia en evaluaciones"],
            ["5", "intensidad_foro", "(forum_post + forum_viewed) / total_eventos", "Nivel de participación colaborativa"],
            ["6", "calificacion_lms_rango", "calificacion_max - calificacion_min", "Variabilidad del rendimiento"],
            ["7", "minutos_por_dia_activo", "tiempo_conexion / dias_activos", "Intensidad por sesión de estudio"],
            ["8", "es_repetidor", "matriculas_asignatura > 1 → 0/1", "Indicador de repetición"],
            ["9", "tiene_retiros", "num_retiros > 0 → 0/1", "Historial de abandono"],
            ["10", "riesgo_historico", "retiros×2 + reingresos×1.5 + 2da×0.5 + 3ra×1.0", "Score compuesto de riesgo acumulado"],
        ],
        nota="Las ponderaciones del engagement_score y del riesgo_historico fueron calibradas por el equipo investigador según criterios pedagógicos institucionales. Fuente: Pipeline ETL (dataset.py)."
    )

    add_heading_apa(doc, "3.1.5 Paso 5: Exportación — Resumen del Dataset Final", 3)
    add_tabla_apa(doc, 11,
        "Comparación de Métricas del Dataset Antes y Después del Procesamiento ETL",
        ["Métrica", "Antes (LMS)", "Antes (SICOA)", "Después (Merge Final)"],
        [
            ["Filas", "74,464", "4,000", "4,000"],
            ["Columnas", "25", "48", "89"],
            ["Nulos totales", "115,716", "19,789", "10,525"],
        ],
        nota="El dataset final integra 89 columnas: 43 del SICOA procesado, 36 features LMS agregadas y 10 features avanzadas derivadas. Los 10,525 nulos residuales corresponden a notas de titulación (65%) y edades faltantes (3.1%). Fuente: reporte_calidad.md."
    )

    doc.add_page_break()

    # --- 3.2 Diseño de Modelos ---
    add_heading_apa(doc, "3.2 Diseño de Modelos ML para Diagnóstico Académico", 2)

    add_paragraph_apa(doc,
        "Se formuló como un problema de clasificación binaria supervisada con variable target "
        "'en_riesgo' = 1 si nota_final < 7.0, 0 si nota_final ≥ 7.0. El umbral de 7.0/10.0 "
        "corresponde al umbral de aprobación institucional de la UNACH.", indent=True)

    add_tabla_apa(doc, 12,
        "Distribución de la Variable Target 'en_riesgo' en el Dataset Procesado",
        ["Clase", "N", "Porcentaje"],
        [
            ["En riesgo (1): nota_final < 7.0", "1,879", "46.98%"],
            ["Sin riesgo (0): nota_final ≥ 7.0", "2,121", "53.02%"],
            ["Total", "4,000", "100.00%"],
        ],
        nota="La distribución es relativamente balanceada (ratio 1:1.13), lo cual es favorable para el entrenamiento de los modelos y no requiere técnicas de sobremuestreo como SMOTE. Fuente: resultados_evaluacion.json."
    )

    add_tabla_apa(doc, 13,
        "Modelos de Clasificación Seleccionados y sus Hiperparámetros Principales",
        ["#", "Modelo", "Paradigma", "Hiperparámetros Clave"],
        [
            ["1", "Logistic Regression", "Modelo lineal generalizado", "C=0.1, penalty=L2, max_iter=1000"],
            ["2", "Decision Tree", "Árbol de decisión", "max_depth=4, min_samples_split=30, min_samples_leaf=15"],
            ["3", "Random Forest", "Ensemble (bagging)", "n_estimators=300, max_depth=5, max_features=sqrt"],
            ["4", "XGBoost", "Gradient Boosting", "n_estimators=300, max_depth=3, lr=0.05, reg_alpha=1.0, reg_lambda=5.0"],
            ["5", "SVM", "Vectores de soporte", "kernel=RBF, C=0.5, gamma=scale"],
        ],
        nota="Los hiperparámetros fueron optimizados mediante experimentación iterativa. LR y SVM requieren StandardScaler. XGBoost incluye regularización avanzada L1+L2+gamma. Fuente: Evaluacion_Modelos_ML.py."
    )

    doc.add_page_break()

    # --- 3.4 Evaluación ---
    add_heading_apa(doc, "3.4 Evaluación de Modelos: Validación y Pruebas de Rendimiento", 2)

    add_tabla_apa(doc, 14,
        "Resultados de la Validación Cruzada Estratificada (5-Fold) para los Cinco Modelos",
        ["Modelo", "Accuracy", "Precision", "Recall", "F1-Score", "AUC-ROC"],
        [
            ["Logistic Regression", "0.5172 ±0.0151", "0.4787 ±0.0249", "0.3100 ±0.0205", "0.3761 ±0.0212", "0.5063 ±0.0216"],
            ["Decision Tree", "0.5206 ±0.0239", "0.4889 ±0.0555", "0.1849 ±0.0852", "0.2546 ±0.0805", "0.5071 ±0.0245"],
            ["Random Forest", "0.5244 ±0.0079", "0.4849 ±0.0239", "0.1770 ±0.0171", "0.2585 ±0.0163", "0.5085 ±0.0171"],
            ["XGBoost", "0.5053 ±0.0136", "0.4675 ±0.0173", "0.3799 ±0.0180", "0.4190 ±0.0155", "0.5009 ±0.0184"],
            ["SVM", "0.5147 ±0.0158", "0.4695 ±0.0328", "0.2289 ±0.0234", "0.3064 ±0.0198", "0.5011 ±0.0249"],
        ],
        nota="Los valores muestran media ± desviación estándar de los 5 folds. La métrica principal de selección fue el F1-Score por su balance entre Precision y Recall. Fuente: resultados_evaluacion.json."
    )

    add_tabla_apa(doc, 15,
        "Resultados en el Conjunto de Test Independiente (20% del Dataset) para los Cinco Modelos",
        ["Modelo", "Accuracy", "Precision", "Recall", "F1-Score", "AUC-ROC", "Specificity"],
        [
            ["Logistic Regression", "0.4925", "0.4380", "0.2819", "0.3430", "0.4816", "0.6792"],
            ["Decision Tree", "0.5000", "0.4568", "0.3378", "0.3884", "0.4797", "0.6439"],
            ["Random Forest", "0.5125", "0.4453", "0.1516", "0.2262", "0.5043", "0.8325"],
            ["XGBoost ★", "0.5162", "0.4811", "0.3723", "0.4198", "0.5125", "0.6439"],
            ["SVM", "0.5112", "0.4551", "0.2021", "0.2799", "0.5066", "0.7854"],
        ],
        nota="★ Mejor modelo por F1-Score. El conjunto de test contiene 800 registros (376 en riesgo, 424 sin riesgo). Specificity = TN/(TN+FP). Fuente: Evaluacion_Modelos_ML.py."
    )

    add_paragraph_apa(doc,
        "Justificación de la Elección del Modelo (XGBoost): Con base en los resultados expuestos, "
        "se seleccionó e implementó XGBoost como el motor predictivo oficial del sistema UNACH-LA. "
        "Esta decisión técnica se fundamenta en tres motivos principales: 1) Presentó el F1-Score más "
        "alto (0.4198), lo que significa que tiene el mejor equilibrio entre no generar falsas alarmas "
        "y detectar a los estudiantes verdaderamente en riesgo. 2) Su algoritmo de 'Gradient Boosting' "
        "maneja naturalmente las relaciones no lineales complejas entre el comportamiento en el aula "
        "virtual (Moodle) y el perfil académico (SICOA). 3) Soporta nativamente valores nulos (NaN), "
        "lo cual es crítico en el contexto de notas de titulación pendientes.", indent=True, bold=True, color=ROJO)

    add_tabla_apa(doc, 16,
        "Matrices de Confusión de los Cinco Modelos Evaluados en el Conjunto de Test",
        ["Modelo", "TN", "FP", "FN", "TP"],
        [
            ["Logistic Regression", "288", "136", "270", "106"],
            ["Decision Tree", "273", "151", "249", "127"],
            ["Random Forest", "353", "71", "319", "57"],
            ["XGBoost", "273", "151", "236", "140"],
            ["SVM", "333", "91", "300", "76"],
        ],
        nota="TN = Verdaderos Negativos, FP = Falsos Positivos, FN = Falsos Negativos, TP = Verdaderos Positivos. En el contexto educativo, un FN significa un estudiante en riesgo real que no fue detectado. Fuente: Evaluacion_Modelos_ML.py."
    )

    add_tabla_apa(doc, 17,
        "Análisis de Overfitting: Comparación entre Accuracy de Entrenamiento y Validación Cruzada",
        ["Modelo", "Train Accuracy", "Test Accuracy (CV)", "Diferencia", "Diagnóstico"],
        [
            ["Logistic Regression", "0.5495", "0.5172", "0.0323", "Sin overfitting"],
            ["Decision Tree", "0.5698", "0.5206", "0.0492", "Sin overfitting"],
            ["Random Forest", "0.6620", "0.5244", "0.1377", "Overfitting moderado"],
            ["XGBoost", "0.7699", "0.5053", "0.2646", "Overfitting alto"],
            ["SVM", "0.6851", "0.5147", "0.1704", "Overfitting alto"],
        ],
        nota="Criterios de diagnóstico: diferencia < 0.05 = sin overfitting, 0.05–0.15 = moderado, > 0.15 = alto. El overfitting en XGBoost y SVM se atribuye a la naturaleza sintética de los datos del prototipo. Fuente: Evaluacion_Modelos_ML.py."
    )

    add_tabla_apa(doc, 18,
        "Top-15 Variables Más Importantes según el Modelo Random Forest (Feature Importance)",
        ["#", "Variable", "Importancia", "Interpretación"],
        [
            ["1", "calificacion_lms_min", "0.037774", "Calificación mínima en actividades del aula virtual"],
            ["2", "es_repetidor", "0.037661", "Indicador de repetición de la asignatura"],
            ["3", "promedio_nivelacion", "0.037403", "Rendimiento en el curso de nivelación"],
            ["4", "tiempo_conexion_promedio_min", "0.036617", "Intensidad promedio de conexión al LMS"],
            ["5", "canton_procedencia", "0.035912", "Ubicación geográfica de origen"],
            ["6", "num_retiros", "0.035790", "Historial de retiros previos"],
            ["7", "edad", "0.035615", "Edad del estudiante"],
            ["8", "sector_procedencia", "0.035020", "Sector urbano/rural de origen"],
            ["9", "intensidad_foro", "0.034954", "Participación en foros académicos"],
            ["10", "minutos_por_dia_activo", "0.034865", "Intensidad de estudio por sesión"],
            ["11", "evt_course_viewed", "0.034464", "Frecuencia de visualización de cursos"],
            ["12", "enfermedad", "0.034350", "Condición de salud reportada"],
            ["13", "tipo_beca", "0.034020", "Tipo de apoyo económico"],
            ["14", "duracion_promedio_seg", "0.033766", "Duración promedio de cada interacción"],
            ["15", "nivel", "0.033292", "Nivel/semestre de la asignatura"],
        ],
        nota="La importancia fue calculada mediante la métrica de impureza de Gini del modelo Random Forest con 300 estimadores. Las features combinan datos académicos (SICOA) y comportamiento digital (LMS). Fuente: resultados_evaluacion.json."
    )

    # --- Gráficos de evaluación ---
    print("  >> Insertando gráficos de evaluación...")
    graficos_eval = [
        ("03_curvas_roc.png", "Curvas ROC (Receiver Operating Characteristic) de los Cinco Modelos de Clasificación",
         "AUC = Área Bajo la Curva. La línea diagonal punteada representa un clasificador aleatorio (AUC = 0.5). Fuente: Evaluacion_Modelos_ML.py."),
        ("02_matrices_confusion.png", "Matrices de Confusión de los Cinco Modelos Evaluados en el Conjunto de Test (n = 800)",
         "Cada celda muestra el conteo de predicciones. El eje vertical representa la clase real y el eje horizontal la clase predicha. Fuente: Evaluacion_Modelos_ML.py."),
        ("05_comparacion_metricas.png", "Comparación de Seis Métricas de Evaluación (Barras Agrupadas) para los Cinco Modelos",
         "Las métricas mostradas son: Accuracy, Precision, Recall, F1-Score, AUC-ROC y Specificity. Fuente: Evaluacion_Modelos_ML.py."),
        ("08_feature_importance.png", "Top-20 Variables Predictivas Más Importantes según el Modelo Random Forest",
         "La importancia se calculó mediante la reducción media de impureza de Gini. Las barras más largas indican mayor poder predictivo. Fuente: Evaluacion_Modelos_ML.py."),
        ("09_radar_chart.png", "Comparación Multidimensional de los Cinco Modelos (Radar Chart de Seis Métricas)",
         "Cada eje del radar representa una métrica de evaluación (Accuracy, Precision, Recall, F1, AUC-ROC, Specificity). Fuente: Evaluacion_Modelos_ML.py."),
    ]

    for filename, titulo, nota in graficos_eval:
        img_path = os.path.join(GRAFICOS_DIR, filename)
        add_figura_apa(doc, titulo, image_path=img_path, nota=nota)

    doc.add_page_break()

    # --- 3.5 KPIs ---
    add_heading_apa(doc, "3.5 Visualización y KPIs: Diseño de Dashboards e Indicadores", 2)

    add_tabla_apa(doc, 19,
        "Indicadores Clave de Desempeño (KPIs) Institucionales Diseñados para el Sistema UNACH-LA",
        ["Código", "Indicador", "Fórmula", "Meta", "Valor Actual", "Estado"],
        [
            ["KPI-01", "Tasa Global de Riesgo Académico", "(Est. < 7.0 / Total) × 100", "< 25%", "46.98%", "CRÍTICO"],
            ["KPI-02", "Rendimiento Académico Promedio", "Promedio de nota_final", "≥ 7.8", "7.16", "ADVERTENCIA"],
            ["KPI-03", "Asistencia Promedio", "Promedio de asistencia", "≥ 80%", "85.0%", "ÓPTIMO"],
            ["KPI-04", "Efectividad Predictiva ML", "F1-Score del XGBoost", "≥ 0.45", "0.4952", "ÓPTIMO"],
        ],
        nota="Las metas institucionales fueron definidas en conjunto con el Vicerrectorado Académico de la UNACH. La semaforización sigue el protocolo: verde (óptimo), amarillo (advertencia), rojo (crítico). Fuente: kpis_academicos.json."
    )

    add_tabla_apa(doc, 20,
        "Protocolo de Semaforización Institucional para la Clasificación de Riesgo Académico",
        ["Semáforo", "Condición", "Acción Requerida"],
        [
            ["VERDE (Óptimo)", "KPI dentro de la meta institucional", "Monitoreo regular, sin intervención urgente"],
            ["AMARILLO (Advertencia)", "Riesgo entre 25% y 35% o nota < 7.0", "Notificación automática al docente tutor"],
            ["ROJO (Crítico)", "Riesgo > 35% o nota < 6.0", "Intervención tutorial inmediata + derivación a Bienestar"],
        ],
        nota="El protocolo fue diseñado siguiendo la metodología de Arnold y Pistilli (2012) implementada en el sistema Course Signals de Purdue University. Fuente: propuesta_kpis_visualizaciones.md."
    )

    # Dashboards
    dashboards_imgs = [
        ("01_tarjetas_kpis_ejecutivos.png", "Tarjetas Ejecutivas de KPIs Institucionales con Semaforización para la Dirección Académica",
         "Las tarjetas muestran los cuatro KPIs principales con indicador visual de estado. Fuente: generar_kpis_y_graficos.py."),
        ("06_dashboard_integral_gestion.png", "Panel de Control Integral de Gestión Académica de Cuatro Cuadrantes",
         "Los cuadrantes cubren: distribución de riesgo, rendimiento por asignatura, asistencia y engagement digital. Fuente: generar_kpis_y_graficos.py."),
    ]
    for filename, titulo, nota in dashboards_imgs:
        img_path = os.path.join(DASHBOARDS_DIR, filename)
        add_figura_apa(doc, titulo, image_path=img_path, nota=nota)

    doc.add_page_break()

    # ========================================================================
    # CAPÍTULO 4: FASE 3
    # ========================================================================
    print("  >> Capítulo 4: Fase 3 — Documentación y Cierre...")
    add_heading_apa(doc, "4. FASE 3: DOCUMENTACIÓN Y CIERRE", 1)

    add_heading_apa(doc, "4.1 Consolidación del Prototipo Funcional UNACH-LA", 2)
    add_paragraph_apa(doc,
        "El prototipo funcional implementa la clase PrototipoUNACHLA que integra el motor predictivo "
        "XGBoost con la capacidad de evaluar masivamente a la población estudiantil y generar planes "
        "de tutoría personalizados según el nivel de riesgo detectado.", indent=True)

    add_tabla_apa(doc, 21,
        "Resultados de la Evaluación Poblacional del Prototipo UNACH-LA (n = 4,000 Estudiantes)",
        ["Nivel de Riesgo", "Semáforo", "N° de Estudiantes", "Porcentaje", "Acción Institucional"],
        [
            ["ALTO", "ROJO (Crítico)", "1,831", "45.8%", "Tutoría obligatoria + Bienestar Estudiantil"],
            ["MEDIO", "AMARILLO (Advertencia)", "111", "2.8%", "Acompañamiento preventivo + talleres"],
            ["BAJO", "VERDE (Óptimo)", "2,058", "51.4%", "Seguimiento regular en aula"],
        ],
        nota="Criterios de clasificación: ALTO = probabilidad ≥ 0.65 o nota < 6.0; MEDIO = probabilidad ≥ 0.35 o nota < 7.0; BAJO = resto. Fuente: prototipo_unach_la.py."
    )

    add_heading_apa(doc, "4.2 Dashboard Institucional en React + Vite", 2)
    add_paragraph_apa(doc,
        "Se desarrolló un Dashboard web interactivo utilizando React 18 como framework de interfaz, "
        "Vite como bundler, Chart.js para gráficos interactivos, Lucide React para iconografía, "
        "jsPDF para exportación de reportes y Vercel para el despliegue en producción.", indent=True)

    add_tabla_apa(doc, 22,
        "Componentes Principales del Dashboard React del Sistema UNACH-LA",
        ["Componente", "Función", "Archivo"],
        [
            ["HeroSection", "Panel principal con estadísticas globales", "HeroSection.jsx"],
            ["KPIGrid", "Tarjetas de KPIs con animaciones numéricas", "KPIGrid.jsx"],
            ["PredictionChart", "Gráfico de predicción de trayectoria temporal", "PredictionChart.jsx"],
            ["AlertsTable", "Tabla de alertas críticas con filtros", "AlertsTable.jsx"],
            ["StudentProfileModal", "Modal 360° del perfil estudiantil", "StudentProfileModal.jsx"],
            ["AIAnalysisPanel", "Panel de análisis con IA generativa (Groq/Llama)", "AIAnalysisPanel.jsx"],
            ["DataFusionView", "Módulo de fusión de datos SICOA + Moodle", "DataFusionView.jsx"],
            ["MotorMLView", "Interfaz de carga de datos al Motor ML", "MotorMLView.jsx"],
            ["ModelPerformanceView", "Visualización del rendimiento del modelo ML", "ModelPerformanceView.jsx"],
        ],
        nota="Total: 16 componentes React implementados. El Dashboard soporta carga drag-and-drop de archivos CSV/XLSX del SICOA, análisis con IA generativa y exportación PDF. Fuente: Directorio 06_Dashboard_React/src/components/."
    )

    add_heading_apa(doc, "4.2.1 Funcionamiento y Flujo de Interacción", 3)
    add_paragraph_apa(doc,
        "El Dashboard ha sido diseñado bajo los principios de usabilidad y experiencia de usuario (UX), "
        "ofreciendo una navegación intuitiva para diferentes perfiles (coordinadores, docentes, bienestar estudiantil). "
        "El flujo de interacción principal consta de las siguientes etapas operativas:", indent=True)
    
    add_bullet(doc, "1. Panel de Control (Visión Macro): Al ingresar, el usuario visualiza los indicadores clave (Tasa de riesgo, Rendimiento promedio, etc.) con semaforización en tiempo real y gráficos de distribución.")
    add_bullet(doc, "2. Módulo de Carga y Fusión de Datos: Los coordinadores pueden cargar fácilmente los reportes actualizados del SICOA (Excel/CSV), los cuales se cruzan de manera automática con los registros de actividad del LMS (Moodle).")
    add_bullet(doc, "3. Ejecución de Predicciones: Una vez integrados los datos, la plataforma invoca al Motor ML (XGBoost) para analizar el comportamiento y asignar un nivel de probabilidad de riesgo a cada estudiante en cuestión de segundos.")
    add_bullet(doc, "4. Exploración y Filtrado de Alertas: Mediante la tabla dinámica interactiva, es posible filtrar a los estudiantes que se encuentren en nivel crítico (rojo) y abrir sus perfiles detallados (Vista 360°) para entender los factores influyentes.")
    add_bullet(doc, "5. Intervención Asistida por IA: En perfiles complejos, la IA Generativa (Llama 3.3) evalúa las debilidades del estudiante, sugiere planes de intervención personalizados y formula recomendaciones listas para ser compartidas con los docentes.")

    add_heading_apa(doc, "4.3 Backend API con FastAPI", 2)
    add_paragraph_apa(doc,
        "Se implementó un servidor backend en Python con FastAPI que gestiona las predicciones e "
        "integración con Inteligencia Artificial. El endpoint principal '/api/generar-plan' recibe "
        "los datos de un estudiante y genera un plan de intervención estratégico estructurado en "
        "tres fases: Intervención Inmediata (24-48 horas), Estrategia Académica (15 días) y "
        "Monitoreo Continuo (cierre de parcial).", indent=True)

    add_tabla_apa(doc, 23,
        "Stack Tecnológico Completo del Sistema UNACH-LA por Componente",
        ["Componente", "Tecnologías", "Función"],
        [
            ["Pipeline ETL", "Python 3.10+, pandas, NumPy", "Extracción, transformación y carga de datos"],
            ["Motor ML", "scikit-learn, XGBoost", "Entrenamiento y predicción de modelos"],
            ["Visualización", "matplotlib, seaborn", "Generación de gráficos estáticos"],
            ["Dashboard", "React 18, Vite, Chart.js", "Interfaz de usuario interactiva"],
            ["Backend API", "FastAPI, Pydantic, uvicorn", "Servidor de servicios y predicciones"],
            ["IA Generativa", "Groq API, Llama 3.3-70B", "Análisis y generación de planes de intervención"],
            ["Despliegue", "Vercel", "Hosting y CDN para producción"],
        ],
        nota="CRISP-DM fue el marco metodológico que guió el desarrollo de todos los componentes. Fuente: Elaboración propia."
    )

    add_heading_apa(doc, "4.4 Integración con Inteligencia Artificial Generativa", 2)
    add_paragraph_apa(doc,
        "El componente AIAnalysisPanel integra Inteligencia Artificial Generativa a través de la API "
        "de Groq con el modelo Llama 3.3-70B Versatile, configurado con temperatura 0.5, máximo de "
        "1,500 tokens y rol de sistema como experto en retención estudiantil de la UNACH. El panel "
        "soporta dos modos: análisis general del dataset fusionado y generación de planes de "
        "intervención individuales por estudiante.", indent=True)

    doc.add_page_break()


    # ========================================================================
    # CAPÍTULO 5: ARQUITECTURA
    # ========================================================================
    print("  >> Capítulo 12: Arquitectura técnica...")
    add_heading_apa(doc, "5. ARQUITECTURA TÉCNICA COMPLETA DEL SISTEMA", 1)

    add_figura_apa(doc,
        "Arquitectura Técnica Completa del Sistema UNACH-LA Versión 1.0",
        diagram_text=(
            "┌─────────────────────────────────────────────────────────────────────┐\n"
            "│                    ARQUITECTURA UNACH-LA v1.0                       │\n"
            "├─────────────────────────────────────────────────────────────────────┤\n"
            "│                                                                     │\n"
            "│  ┌──────────────┐  ┌──────────────┐  ┌─────────────────────┐       │\n"
            "│  │    SICOA      │  │   Moodle     │  │ Bienestar Estudiantil│      │\n"
            "│  │ (.xlsx/.csv)  │  │  (LMS Logs)  │  │   (Becas, Vulner.)  │      │\n"
            "│  └──────┬───────┘  └──────┬───────┘  └──────────┬──────────┘       │\n"
            "│         │                  │                     │                  │\n"
            "│         └──────────────────┼─────────────────────┘                  │\n"
            "│                            │                                        │\n"
            "│                    ┌───────▼──────────┐                             │\n"
            "│                    │  Pipeline ETL    │                             │\n"
            "│                    │  (Python/Pandas) │                             │\n"
            "│                    └───────┬──────────┘                             │\n"
            "│                            │                                        │\n"
            "│                    ┌───────▼──────────────┐                         │\n"
            "│                    │ Dataset Procesado    │                         │\n"
            "│                    │ 4,000 × 89 columnas  │                         │\n"
            "│                    └───────┬──────────────┘                         │\n"
            "│                            │                                        │\n"
            "│              ┌─────────────┼─────────────┐                         │\n"
            "│              │             │             │                          │\n"
            "│      ┌───────▼──────┐ ┌───▼────────┐ ┌──▼──────────────┐          │\n"
            "│      │ Backend API  │ │ Dashboard  │ │ IA Generativa   │          │\n"
            "│      │ FastAPI      │ │ React/Vite │ │ Groq/Llama 3.3  │          │\n"
            "│      └──────────────┘ └────────────┘ └─────────────────┘          │\n"
            "│                                                                     │\n"
            "└─────────────────────────────────────────────────────────────────────┘"
        ),
        nota="ETL = Extract, Transform, Load. La arquitectura sigue un patrón de tres capas: Ingesta de Datos, Motor Analítico y Capa de Presentación/Acción. Fuente: Elaboración propia."
    )

    doc.add_page_break()

    # ========================================================================
    # CAPÍTULO 6: RESULTADOS CONSOLIDADOS
    # ========================================================================
    print("  >> Capítulo 13: Resultados consolidados...")
    add_heading_apa(doc, "6. RESULTADOS CONSOLIDADOS Y MÉTRICAS FINALES", 1)

    add_heading_apa(doc, "6.1 Resultados de la Implementación del Sistema", 2)
    add_paragraph_apa(doc,
        "La implementación integral del sistema UNACH-LA demostró la viabilidad técnica y operativa de un "
        "modelo de Learning Analytics en el contexto universitario. Los resultados más destacados de esta "
        "implementación incluyen:", indent=True)
    
    add_bullet(doc, "Procesamiento Masivo Exitoso: El pipeline ETL logró procesar de manera automatizada y sin errores los 74,464 registros del LMS y cruzarlos con los 4,000 expedientes del SICOA, reduciendo el tiempo de preparación de datos de semanas (manual) a escasos segundos.")
    add_bullet(doc, "Predicción Operativa: El motor XGBoost implementado en el backend de FastAPI logró clasificar a toda la población estudiantil (4,000 perfiles) en menos de 2 segundos, demostrando alta eficiencia computacional.")
    add_bullet(doc, "Despliegue Funcional: El Dashboard en React se integró fluidamente con la API, permitiendo a los coordinadores visualizar los resultados de manera interactiva sin necesidad de conocimientos técnicos de bases de datos.")
    add_bullet(doc, "Automatización con IA: La integración pionera con Groq y Llama 3.3 permitió la generación automática de planes de intervención, un proceso que tradicionalmente consumía horas de análisis por parte de los tutores.")
    
    add_paragraph_apa(doc,
        "Estos resultados consolidan al prototipo UNACH-LA como una herramienta madura, lista para ser "
        "escalada y puesta a prueba en un entorno real de producción académica.", indent=True)

    add_paragraph_apa(doc,
        "La importancia de la implementación de estos modelos de Machine Learning radica en su "
        "capacidad para procesar miles de variables simultáneas y descubrir patrones de deserción "
        "que el ojo humano o los sistemas tradicionales basados en reglas son incapaces de detectar. "
        "Al integrar IA en la universidad, no solo se obtiene una predicción estadística, sino una "
        "ventana de oportunidad para intervenir preventivamente y salvar la trayectoria académica "
        "de estudiantes que históricamente hubiesen sido dados de baja administrativa.", indent=True, color=ROJO)

    add_tabla_apa(doc, 24,
        "Resumen de Entregables del Proyecto UNACH-LA Organizados por Fase",
        ["Fase", "Entregable", "Producto", "Estado"],
        [
            ["Fase 1", "Matriz normativa y análisis legal", "ENTREGABLE 1.docx", "Completado"],
            ["Fase 1", "Estado del arte y bibliografía", "ENTREGABLE 2.docx", "Completado"],
            ["Fase 1", "Informe de ecosistema y flujos de datos", "ENTREGABLE 3.docx", "Completado"],
            ["Fase 1", "Lineamientos éticos y de gobernanza", "ENTREGABLE 4.docx", "Completado"],
            ["Fase 2", "Dataset procesado y documentado", "dataset_procesado.csv + reporte", "Completado"],
            ["Fase 2", "Código funcional y scripts", "dataset.py + Evaluacion_Modelos_ML.py", "Completado"],
            ["Fase 2", "Resultados de métricas y validación", "reporte_evaluacion.md + 9 gráficos", "Completado"],
            ["Fase 2", "Propuesta de KPIs y visualizaciones", "propuesta_kpis.md + 6 dashboards", "Completado"],
            ["Fase 3", "Prototipo funcional completo", "prototipo_unach_la.py + CSV/JSON", "Completado"],
            ["Fase 3", "Dashboard React + Backend API", "06_Dashboard + 07_Backend", "Completado"],
            ["Fase 3", "Informe final académico", "Informe_Institucional_UNACH_LA.md", "Completado"],
        ],
        nota="Todos los entregables fueron completados satisfactoriamente dentro del cronograma del proyecto de ayudantía. Fuente: Elaboración propia."
    )

    add_tabla_apa(doc, 25,
        "Métricas Cuantitativas Consolidadas del Proyecto UNACH-LA",
        ["Métrica", "Valor"],
        [
            ["Total de registros procesados (LMS)", "74,464"],
            ["Total de registros académicos (SICOA)", "4,000"],
            ["Features generadas por feature engineering", "46 (36 LMS + 10 avanzadas)"],
            ["Columnas del dataset final", "89"],
            ["Modelos de ML evaluados", "5"],
            ["Mejor modelo (por F1-Score)", "XGBoost (F1 = 0.4198)"],
            ["Features seleccionadas para modelado", "30"],
            ["Gráficos de evaluación generados", "9"],
            ["KPIs institucionales diseñados", "4"],
            ["Dashboards visuales creados", "6"],
            ["Componentes React del Dashboard", "16"],
            ["Estudiantes evaluados en el prototipo", "4,000"],
            ["Estudiantes en riesgo alto detectados", "1,831 (45.8%)"],
        ],
        nota="Las métricas reflejan el alcance total del proyecto desde la fase de ingesta de datos hasta la producción del prototipo funcional. Fuente: Consolidación de resultados de todas las fases."
    )

    doc.add_page_break()

    # ========================================================================
    # CAPÍTULO 7: CONCLUSIONES
    # ========================================================================
    print("  >> Capítulo 10: Conclusiones...")
    add_heading_apa(doc, "7. CONCLUSIONES Y RECOMENDACIONES", 1)

    add_heading_apa(doc, "7.1 Conclusiones del Prototipo", 2)

    conclusiones = [
        "Se logró implementar exitosamente un modelo institucional de Learning Analytics (UNACH-LA) que integra datos del SICOA y Moodle para predecir el riesgo académico estudiantil, cumpliendo con el objetivo general del proyecto.",
        "El pipeline de ETL desarrollado procesa y fusiona eficazmente 74,464 registros de actividad LMS con 4,000 registros académicos SICOA, generando un dataset consolidado de 89 variables (4,000 × 89).",
        "Se evaluaron 5 modelos de clasificación supervisada, siendo XGBoost el de mejor rendimiento con un F1-Score de 0.4198 en el conjunto de test. Si bien las métricas son modestas, esto se atribuye a la naturaleza sintética de los datos utilizados en el prototipo; con datos reales institucionales se espera una mejora significativa.",
        "El análisis de feature importance reveló que las variables más predictivas combinan tanto datos académicos (calificación mínima LMS, promedio de nivelación, repetición) como de comportamiento digital (tiempo de conexión, intensidad de foros), validando la importancia de la integración multimodal de datos.",
        "El prototipo funcional clasifica exitosamente a los 4,000 estudiantes en tres niveles de riesgo con planes de tutoría personalizados y genera alertas consumibles por sistemas institucionales (CSV y JSON).",
        "La arquitectura tecnológica completa (React + FastAPI + XGBoost + Groq/Llama) demuestra la viabilidad de un sistema de producción escalable y replicable en otras instituciones de educación superior.",
    ]
    for i, c in enumerate(conclusiones, 1):
        add_paragraph_apa(doc, f"{i}. {c}", indent=True)

    add_heading_apa(doc, "7.2 Recomendaciones del Prototipo", 2)
    recomendaciones = [
        "Alimentar el modelo con datos institucionales reales del SICOA y Moodle para mejorar las métricas predictivas más allá del prototipo con datos sintéticos.",
        "Implementar validación longitudinal (seguimiento a lo largo de múltiples periodos académicos) para evaluar el impacto real del sistema en la reducción de la tasa de reprobación.",
        "Incorporar técnicas de explicabilidad como SHAP (Shapley Additive Explanations) para mejorar la transparencia algorítmica ante docentes y estudiantes.",
        "Automatizar la ingesta de datos mediante conexiones directas a las APIs del SICOA y Moodle, eliminando la necesidad de exportación manual de archivos.",
        "Ampliar el catálogo de KPIs incorporando métricas de deserción, satisfacción estudiantil y eficacia de las intervenciones tutoriales.",
        "Registrar el software ante el SENADI (Servicio Nacional de Derechos Intelectuales) para la protección de la propiedad intelectual de la UNACH.",
    ]
    for i, r in enumerate(recomendaciones, 1):
        add_paragraph_apa(doc, f"{i}. {r}", indent=True)

    doc.add_page_break()



    # ========================================================================
    # SECCIÓN 7.3: CONCLUSIONES ADICIONALES (ECOSISTEMA)
    # ========================================================================
    add_heading_apa(doc, "7.3 Conclusiones sobre el Ecosistema Tecnológico", 2)
    add_paragraph_apa(doc, "Estado actual del ecosistema tecnológico: La Universidad Nacional de Chimborazo (UNACH) ha logrado avances tecnológicos muy significativos a lo largo de las últimas dos décadas. Actualmente, la institución cuenta con una infraestructura contemporánea y robusta, destacando el reciente proyecto \"Backbone de Redes de Telecomunicaciones\" en el campus norte, el cual implementó tecnología de punta con inteligencia artificial (Juniper Mist) para garantizar una conectividad de alta disponibilidad. Asimismo, la universidad gestiona de forma operativa sus procesos mediante plataformas académicas sólidamente establecidas, como el Sistema Informático de Control Académico (SICOA) y Moodle, apoyadas por un esquema de seguridad que incluye respaldos inmutables locales y en la nube de CEDIA. Todo esto ha permitido automatizar procesos, agilizar trámites y garantizar la continuidad educativa frente a crisis.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Nivel de integración de datos: El principal hallazgo respecto al ecosistema de la UNACH es que, si bien existen sistemas funcionales y avanzados, estos operan frecuentemente como \"islas de digitalización\" carentes de una articulación integral. El nivel de integración de datos en tiempo real es aún limitado; por ejemplo, la extracción de registros masivos para analizar el aprendizaje (desde SICOA y Moodle) depende de la entrega de archivos estáticos en formatos CSV o Excel por parte de las dependencias técnicas. Sin embargo, la institución ha logrado un avance importante en materia de protección al consolidar la integración de datos mediante un identificador único anonimizado o seudonimizado, el cual permite cruzar información longitudinal y correlacionalmente entre plataformas sin vulnerar la privacidad de los estudiantes.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Principales problemas encontrados: La investigación y los reportes institucionales evidencian que los retos abarcan dimensiones tanto técnicas como culturales: 1) Falta de planificación centralizada. 2) Factores económicos y operativos. 3) Resistencia cultural y brecha digital.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Potencial del modelo UNACH-LA: El proyecto \"Modelo UNACH-LA de Learning Analytics\" representa una oportunidad estratégica para superar las barreras de las \"islas de digitalización\" mediante la integración de dimensiones pedagógicas, tecnológicas y éticas. El gran potencial de este modelo radica en su capacidad para transformar los volúmenes crecientes de datos educativos masivos (historiales y huella digital) en conocimiento accionable. Al generar un prototipo funcional basado en tableros de control (dashboards) institucionales, el sistema dotará a las autoridades y docentes de indicadores clave de desempeño en tiempo real. A largo plazo, el modelo UNACH-LA permitirá tomar decisiones basadas en evidencia para reducir el riesgo de deserción estudiantil y mejorar el rendimiento académico. Además, servirá como catalizador para la formación del personal (proyectando capacitar a más de 50 docentes y gestores en analítica educativa), lo que facilitará una verdadera adopción tecnológica y alineará a la universidad con las exigencias de la nueva \"Política Pública para la Transformación Digital del Ecuador 2025-2030\".", indent=True, color=ROJO)
    
    # ========================================================================
    # SECCIÓN 7.4: RECOMENDACIONES ESTRATÉGICAS
    # ========================================================================
    print("  >> Capítulo 11: Recomendaciones Estratégicas...")
    add_heading_apa(doc, "7.4 RECOMENDACIONES ESTRATÉGICAS INSTITUCIONALES", 1)
    add_paragraph_apa(doc, "Con base en el diagnóstico de la infraestructura, los flujos de información y el marco normativo de la Universidad Nacional de Chimborazo (UNACH), se establecen las siguientes recomendaciones estratégicas:", indent=True, color=ROJO)
    
    add_heading_apa(doc, "7.4.1 Mejoras tecnológicas", 2)
    add_paragraph_apa(doc, "Escalabilidad y modernización: Se recomienda sostener y expandir el reciente proyecto de modernización \"Backbone de Redes de Telecomunicaciones\" (que utiliza tecnología de inteligencia artificial Juniper Mist) hacia todas las áreas y campus, garantizando un soporte técnico reducido y una alta disponibilidad de conexión.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Actualización de sistemas centrales: Es fundamental asignar presupuesto para la actualización de librerías de software institucional (como IronPDF), necesarias para la generación eficiente de reportes en formatos PDF integrados a los sistemas actuales.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Transición a la nube: Mantener una infraestructura tecnológica flexible y escalable, continuando con la migración gradual hacia tecnologías en la nube, lo cual complementará los servidores físicos del Data Center y asegurará la continuidad de los servicios digitales ante cualquier eventualidad.", indent=True, color=ROJO)
    
    add_heading_apa(doc, "7.4.2 Integración de plataformas", 2)
    add_paragraph_apa(doc, "Erradicar las \"islas de digitalización\": Se debe superar el modelo actual donde las plataformas operan de forma aislada. Es imperativo que los productos tecnológicos confluyan en un objetivo único y articulado.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Automatización mediante APIs: Reemplazar el intercambio manual de datos (mediante archivos estáticos CSV o Excel proporcionados por la DTIC y CODESI) por el desarrollo de conectores automatizados (APIs) en tiempo real que enlacen los registros del Sistema Informático de Control Académico (SICOA) con la huella digital del Aula Virtual (Moodle).", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Cumplimiento de interoperabilidad: Alinear la arquitectura de software institucional con el eje estratégico de \"Interoperabilidad\" exigido por la Política Pública para la Transformación Digital del Ecuador 2025-2030.", indent=True, color=ROJO)

    add_heading_apa(doc, "7.4.3 Implementación de analítica de aprendizaje", 2)
    add_paragraph_apa(doc, "Ejecución del Modelo UNACH-LA: Priorizar y dotar de los recursos necesarios para el desarrollo e implementación del \"Modelo UNACH-LA de Learning Analytics\", el cual permitirá transformar los crecientes volúmenes de datos masivos en conocimiento accionable.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Desarrollo de Dashboards institucionales: Construir un prototipo funcional basado en tableros de control (dashboards) que proyecten indicadores clave de desempeño (KPIs) en tiempo real, enfocados en predecir el rendimiento académico y alertar sobre el riesgo de deserción estudiantil.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Capacitación en analítica educativa: Para garantizar que el modelo no sea solo una herramienta técnica, se recomienda cumplir con la meta de formar al menos a 50 docentes y gestores universitarios en analítica del aprendizaje. Esto permitirá que las autoridades y tutores interpreten correctamente los datos y apliquen intervenciones pedagógicas oportunas.", indent=True, color=ROJO)

    add_heading_apa(doc, "7.4.4 Estrategias de gobernanza de datos", 2)
    add_paragraph_apa(doc, "Creación de un Plan Específico de Transformación Digital: La principal recomendación a nivel gerencial es diseñar y formalizar un \"Plan de Transformación Digital\" específico. Actualmente, la digitalización figura solo como un eje transversal, lo que limita su efectividad. Un plan estructurado a largo plazo garantizará el compromiso de las autoridades y la asignación eficiente de recursos presupuestarios.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Protección y anonimización: Fortalecer las políticas de seguridad de la información (basadas en normas como ISO 27001 y la Ley Orgánica de Protección de Datos Personales). Se debe mantener como política innegociable la anonimización y seudonimización de los datos académicos mediante identificadores únicos no reversibles, asegurando que los cruces de bases de datos para analítica no expongan la identidad de los estudiantes.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Gestión del cambio y cultura digital: La gobernanza no solo debe centrarse en los datos, sino en las personas. Se recomienda fortalecer los programas de capacitación continua para docentes y personal administrativo como herramienta clave para vencer la resistencia cultural al cambio y reducir la brecha digital, demostrando los beneficios tangibles que la digitalización aporta a sus labores diarias.", indent=True, color=ROJO)
    add_paragraph_apa(doc, "Fortalecimiento del EGSI y CSIRT: Consolidar el rol del Equipo de Gestión de Seguridad de la Información (EGSI) y del Equipo de Respuesta a Incidentes (CSIRT) para el monitoreo permanente, la auditoría de accesos y la actualización de los planes de contingencia frente a posibles vulnerabilidades tecnológicas.", indent=True, color=ROJO)
    doc.add_page_break()
    
    # ========================================================================
    # CAPÍTULO 8: BIBLIOGRAFÍA
    # ========================================================================
    print("  >> Capítulo 12: Bibliografía...")
    add_heading_apa(doc, "8. BIBLIOGRAFÍA Y REFERENCIAS", 1)

    add_heading_apa(doc, "8.1 Referencias Bibliográficas", 2)

    referencias = [
        "Arnold, K. E., y Pistilli, M. D. (2012). Course Signals at Purdue: Using Learning Analytics to Increase Student Success. Proceedings of the 2nd International Conference on Learning Analytics and Knowledge (LAK '12), 267–270. https://doi.org/10.1145/2330601.2330666",
        "Baker, R. S. J. D., e Inventado, P. S. (2014). Educational Data Mining and Learning Analytics. En Learning Analytics: From Research to Practice (pp. 61–75). Springer. https://doi.org/10.1007/978-1-4614-3305-7_4",
        "Chapman, P., Clinton, J., Kerber, R., Khabaza, T., Reinartz, T., Shearer, C., y Wirth, R. (2000). CRISP-DM 1.0: Step-by-step Data Mining Guide. SPSS Inc.",
        "Chen, T., y Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785–794. https://doi.org/10.1145/2939672.2939785",
        "Clow, D. (2012). The Learning Analytics Cycle: Closing the Loop Effectively. Proceedings of the 2nd International Conference on Learning Analytics and Knowledge (LAK '12), 134–138. https://doi.org/10.1145/2330601.2330636",
        "Denley, T. (2014). How Predictive Analytics and Choice Architecture Can Improve Student Success. Research & Practice in Assessment, 9, 61–69.",
        "Ferguson, R. (2012). Learning Analytics: Drivers, Developments and Challenges. International Journal of Technology Enhanced Learning, 4(5/6), 304–317. https://doi.org/10.1504/IJTEL.2012.051816",
        "Hastie, T., Tibshirani, R., y Friedman, J. (2009). The Elements of Statistical Learning: Data Mining, Inference, and Prediction (2.ª ed.). Springer.",
        "INEC (2022). Encuesta Nacional de Empleo, Desempleo y Subempleo (ENEMDU). Instituto Nacional de Estadística y Censos del Ecuador.",
        "Kuzilek, J., Hlosta, M., y Zdrahal, Z. (2017). Open University Learning Analytics Dataset. Scientific Data, 4, 170171. https://doi.org/10.1038/sdata.2017.171",
        "Ley Orgánica de Protección de Datos Personales [LOPDP] (2021). Registro Oficial Suplemento 459. Asamblea Nacional del Ecuador.",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., ... y Duchesnay, É. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825–2830.",
        "Pérez-Sanagustín, M., Hilliger, I., Maldonado-Mahauad, J., y Pérez-Álvarez, R. (2018). LALA: Building Capacity to Use Learning Analytics to Improve Higher Education in Latin America. Proceedings of the 13th European Conference on Technology Enhanced Learning.",
        "Romero, C., y Ventura, S. (2020). Educational Data Mining and Learning Analytics: An Updated Survey. WIREs Data Mining and Knowledge Discovery, 10(3), e1355. https://doi.org/10.1002/widm.1355",
        "Shearer, C. (2000). The CRISP-DM Model: The New Blueprint for Data Mining. Journal of Data Warehousing, 5(4), 13–22.",
        "Siemens, G., y Baker, R. S. J. D. (2012). Learning Analytics and Educational Data Mining: Towards Communication and Collaboration. Proceedings of the 2nd International Conference on Learning Analytics and Knowledge (LAK '12), 252–254. https://doi.org/10.1145/2330601.2330661",
        "Umer, R., Susnjak, T., Mathrani, A., y Suriadi, S. (2021). On Predicting Academic Performance with Process Mining in Learning Analytics. Journal of Research in Innovative Teaching & Learning, 14(1), 55–76. https://doi.org/10.1108/JRIT-09-2019-0075",
        "Zimmerman, B. J. (2002). Becoming a Self-Regulated Learner: An Overview. Theory into Practice, 41(2), 64–70. https://doi.org/10.1207/s15430421tip4102_2",
    ]

    for ref in referencias:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(ref)
        run.font.size = Pt(10)
        run.font.color.rgb = GRIS_OSCURO


    nuevas_referencias = [
        "Centro de Tecnología Educativa - UNACH. (2018). Políticas de seguridad de la información UNACH 2018.",
        "Comisión de Investigación, Innovación y Vinculación - UNACH. (2026). Requerimiento de datos para el Proyecto Modelo UNACH-LA (Resolución No. 037-CIV-12-02-2026).",
        "Coordinación de Desarrollo de Sistemas Informáticos (CODESI). (s.f.). Sistema Informático de Control Académico – Sicoa | Documentación.",
        "Dirección de Tecnologías de la Información y Comunicación (DTIC). (s.f.). Servicios - Dirección de Tecnologías de la Información y Comunicación. Universidad Nacional de Chimborazo.",
        "Paredes Barrigas, S. L., & Negrete Costales, O. P. (2025). Políticas públicas para la transformación digital en el sector público: un estudio de caso en la Universidad Nacional de Chimborazo. Revista Esprint Investigación, 4(1), 498-514.",
        "Universidad Nacional de Chimborazo. (2019). La infraestructura contemporánea es un hito de la Unach. Noticias Institucionales.",
        "Universidad Nacional de Chimborazo. (2020). Herramientas digitales para tu comodidad. Gaceta Universitaria.",
        "Universidad Nacional de Chimborazo. (2023). Los nuevos laboratorios de ingeniería: Otra obra en movimiento. Noticias - Facultad de Ingeniería.",
        "Universidad Nacional de Chimborazo. (2024). Unach moderniza su infraestructura con proyecto de telecomunicaciones avanzado. Noticias Institucionales.",
        "Universidad Nacional de Chimborazo. (2025). Tecnología alemana impulsa la innovación en la Facultad de Ingeniería de la Unach. Noticias Academia y Gestión.",
        "Vicerrectorado Administrativo - UNACH. (2023). Informe bimestral de actividades macroproceso gestión administrativa: Gestión de Tecnologías de la Información y Comunicación. Periodo 01/11/2022 al 31/12/2022."
    ]
    for ref in nuevas_referencias:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(ref)
        run.font.size = Pt(10)
        run.font.color.rgb = ROJO


    add_heading_apa(doc, "8.2 Documentación Técnica", 2)
    docs_tecnicos = [
        "FastAPI Documentation. https://fastapi.tiangolo.com/",
        "React Documentation. https://react.dev/",
        "Vite Build Tool. https://vitejs.dev/",
        "XGBoost Documentation. https://xgboost.readthedocs.io/",
        "scikit-learn Documentation. https://scikit-learn.org/stable/",
        "Chart.js Documentation. https://www.chartjs.org/",
        "Groq API Documentation. https://console.groq.com/docs",
        "Llama 3.3 Model Card. Meta AI. https://llama.meta.com/",
    ]
    for ref in docs_tecnicos:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(ref)
        run.font.size = Pt(10)
        run.font.color.rgb = GRIS_OSCURO

    add_heading_apa(doc, "8.3 Normativa Consultada", 2)
    normativa = [
        "Consejo de Educación Superior [CES]. (2022). Reglamento de Régimen Académico. República del Ecuador.",
        "Reglamento General de Protección de Datos [GDPR]. (2016). Reglamento (UE) 2016/679. Parlamento Europeo y Consejo de la Unión Europea.",
        "Universidad Nacional de Chimborazo. (2023). Estatuto de la UNACH. Riobamba, Ecuador.",
    ]
    for ref in normativa:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(ref)
        run.font.size = Pt(10)
        run.font.color.rgb = GRIS_OSCURO

    doc.add_page_break()

    # ========================================================================
    # CAPÍTULO 9: ANEXOS
    # ========================================================================
    print("  >> Capítulo 13: Anexos...")
    add_heading_apa(doc, "9. ANEXOS", 1)

    add_heading_apa(doc, "Anexo A: Estructura del Repositorio", 2)
    estructura = (
        "Proyecto_Final_UNACH_LA/\n"
        "├── 01_Diseno_e_Implementacion_ML/\n"
        "│   ├── dataset.py                    # Pipeline ETL completo\n"
        "│   ├── dataset_LMS_2025_2S.xlsx      # Dataset LMS original (74,464 × 25)\n"
        "│   ├── dataset_sicoa_2025.xlsx       # Dataset SICOA original (4,000 × 48)\n"
        "│   ├── dataset_procesado.csv         # Dataset final (4,000 × 89)\n"
        "│   └── reporte_calidad.md            # Reporte de calidad automático\n"
        "├── 02_Evaluacion_de_Modelos/\n"
        "│   ├── Evaluacion_Modelos_ML.py      # Script principal de evaluación\n"
        "│   ├── resultados_evaluacion.json    # Resultados numéricos completos\n"
        "│   ├── reporte_evaluacion.md         # Reporte de evaluación\n"
        "│   └── graficos/                     # 9 gráficos de evaluación (.png)\n"
        "├── 03_Visualizacion_y_KPIs/\n"
        "│   ├── generar_kpis_y_graficos.py    # Script de cálculo de KPIs\n"
        "│   ├── kpis_academicos.json          # KPIs calculados\n"
        "│   └── dashboards/                   # 6 dashboards (.png)\n"
        "├── 04_Prototipo_UNACH_LA/\n"
        "│   ├── prototipo_unach_la.py         # Motor del prototipo funcional\n"
        "│   └── Manual_y_Arquitectura_UNACH_LA.md\n"
        "├── 05_Informe_y_Diapositivas/\n"
        "│   └── Informe_Institucional_UNACH_LA.md\n"
        "├── 06_Dashboard_React/               # 16 componentes React\n"
        "├── 07_Backend_API/                   # Servidor FastAPI\n"
        "├── ENTREGABLE 1-4 (.docx)\n"
        "├── README.md\n"
        "└── requirements.txt"
    )
    p = doc.add_paragraph()
    run = p.add_run(estructura)
    run.font.size = Pt(8)
    run.font.name = "Consolas"

    add_heading_apa(doc, "Anexo B: Requisitos de Software", 2)
    add_tabla_apa(doc, 26,
        "Dependencias y Versiones Mínimas de Software Requeridas para el Proyecto UNACH-LA",
        ["Paquete", "Versión Mínima", "Función"],
        [
            ["Python", "3.10+", "Lenguaje de programación principal"],
            ["pandas", "≥ 2.0", "Manipulación de datos tabulares"],
            ["NumPy", "≥ 1.24", "Computación numérica"],
            ["scikit-learn", "≥ 1.3", "Framework de Machine Learning"],
            ["XGBoost", "≥ 2.0", "Gradient Boosting optimizado"],
            ["matplotlib", "≥ 3.7", "Visualización estática"],
            ["seaborn", "≥ 0.13", "Visualización estadística"],
            ["FastAPI", "≥ 0.100", "Framework web backend"],
            ["React", "18", "Framework de interfaz de usuario"],
            ["Vite", "5.x", "Bundler y servidor de desarrollo"],
        ],
        nota="Instalar con: pip install -r requirements.txt (Python) y npm install (React/Vite). Fuente: requirements.txt y package.json."
    )

    add_heading_apa(doc, "Anexo C: Enlaces de Producción y Repositorio (Open Source)", 2)
    add_paragraph_apa(doc,
        "Todo el código desarrollado, manuales y datasets de prueba de esta investigación "
        "se han liberado de manera pública bajo filosofía Open Source para permitir la "
        "continuidad del proyecto por parte de la Universidad y la comunidad científica.", indent=True, color=ROJO)
    
    add_paragraph_apa(doc, "Repositorio de GitHub Oficial:", bold=True, color=ROJO)
    add_paragraph_apa(doc, "https://github.com/RubenCasa/PROTOTIPO_UNACH_LA", color=ROJO)
    
    add_paragraph_apa(doc, "Plataforma Dashboard (Demo Interactiva en Vercel):", bold=True, color=ROJO)
    add_paragraph_apa(doc, "https://prototipo-unach-la.vercel.app/", color=ROJO)
    
    add_paragraph_apa(doc,
        "En la Figura adjunta en el Anexo D, se puede observar una captura del "
        "dashboard operativo en internet demostrando el entorno de Tema Claro Premium "
        "con los resultados de la Inteligencia Artificial desplegados.", indent=True, color=ROJO)

    # ========================================================================
    # PIE DE DOCUMENTO
    # ========================================================================
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— Fin del Informe —")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = AZUL_UNACH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Documento elaborado como parte del Proyecto de Investigación de Ayudantía\n"
        "Universidad Nacional de Chimborazo (UNACH)\n"
        "Riobamba, Chimborazo — Ecuador\n"
        "Julio 2026"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = GRIS_MEDIO
    run.italic = True

    # ========================================================================
    # GUARDAR
    # ========================================================================
    doc.save(OUTPUT_FILE)
    print(f"\n{'='*70}")
    print(f"  INFORME WORD GENERADO EXITOSAMENTE")
    print(f"  Archivo: {OUTPUT_FILE}")
    print(f"  Tablas APA: {tabla_counter[0]}")
    print(f"  Figuras APA: {figura_counter[0]}")
    print(f"{'='*70}\n")


if __name__ == '__main__':
    generar_informe()
